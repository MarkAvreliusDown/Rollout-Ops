
import calendar
import copy
import datetime
import ipaddress
import json
import logging
import re

import requests
from django.conf import settings
from django.db.models import Q
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone

from django.views.decorators.http import require_POST

from .icons import VALID_COLORS, VALID_ICONS
from .models import (
    ANCHOR_CHOICES, AREA_FORMAT_CHOICES, BRANCH_CHOICES, COMPANY_INN, CONTRACTOR_CHOICES, DEVICE_DNS1, DEVICE_DNS2,
    DEVICE_IP_LABELS, DEVICE_MASK, KSO_NUMBER_BASE, STORE_KANBAN_STATUS_CHOICES,
    STORE_TYPE_CHOICES, Article,
    ArticleFile, ArticleImage, Bonus, BudgetCalculation, BudgetSettings, ChecklistItem, Contact, Contractor,
    KbCategory, KbCategoryFile, LetterCompany, Note, Notification, NoteFile, Problem, ReconstructionDocument, ReconstructionRecord, Report,
    ReportPhoto, RouterConfigTemplate, SalarySettings, Store, StoreCashRegister, StoreDeviceIPConfig, StoreLog, StoreLogFile,
    StoreRouterConfig, Task,
    compute_budget_totals, compute_device_ips_for_network, compute_kso_ips, compute_router_config_for_store,
    sync_new_store_to_letters,
)

MONTH_NAMES_RU = [
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
]

# Подписи/подсказки для плейсхолдеров шаблона конфига роутера — см. инструкцию
# по конфигурации FortiGate. Для плейсхолдеров, которых здесь нет (например, если
# шаблон обновят и добавят новые переменные), используется само имя плейсхолдера.
ROUTER_CONFIG_LABELS = {
    "SHOPID": ("Номер магазина", "пример: 1481"),
    "WANIP": ("IP адрес провайдера 1", "пример: 192.0.2.10"),
    "WANMASK": ("IP маска провайдера 1", "пример: 255.255.255.252"),
    "TUNIP": ("IP адрес туннеля провайдера 1", "встречается 2 раза в конфиге — подставится в оба места"),
    "WANGW": ("IP шлюз провайдера 1", "пример: 192.0.2.9"),
    "WANIP2": ("IP адрес провайдера 2", "пример: 198.51.100.10"),
    "WANMASK2": ("IP маска провайдера 2", "пример: 255.255.255.248"),
    "TUNIP2": ("IP адрес туннеля провайдера 2", ""),
    "WANGW2": ("IP шлюз провайдера 2", ""),
    "NETWORK": ("Подсеть LAN (без маски)", "пример: 10.90.240.0"),
    "LANIP": ("IP адрес LAN сети", "пример: 10.90.240.1, встречается >1 раза в конфиге"),
    "LANMASK": ("IP маска LAN сети", "пример: 255.255.255.128"),
    "DHCPSTART": ("Первый адрес DHCP пула", "подставляется автоматически по LANIP, можно поправить вручную"),
    "DHCPEND": ("Последний адрес DHCP пула", "подставляется автоматически по LANIP, можно поправить вручную"),
}


def _router_config_placeholders(store):
    """Список плейсхолдеров <NAME> из текущего шаблона конфига роутера (в порядке
    первого появления, без повторов) вместе с сохранёнными для магазина значениями."""
    template_text = RouterConfigTemplate.get().template_text
    keys = list(dict.fromkeys(re.findall(r"<([A-Z0-9_]+)>", template_text)))
    saved = StoreRouterConfig.objects.filter(store=store).first()
    values = dict(saved.values) if saved else {}
    if not values.get("SHOPID") and store.number:
        values["SHOPID"] = store.number
    placeholders = []
    for key in keys:
        label, hint = ROUTER_CONFIG_LABELS.get(key, (key, ""))
        placeholders.append({"key": key, "label": label, "hint": hint, "value": values.get(key, "")})
    return placeholders, keys

logger = logging.getLogger("board")


def _active_llm_name():
    return settings.OLLAMA_MODEL


def _llm_ready():
    return True


def _llm_not_ready_message():
    return "Ollama недоступна (проверьте, что приложение Ollama запущено)"


def _log(log_list, message, question=None):
    log_list.append(message)
    if question:
        logger.info("[вопрос: %s] %s", question, message)
    else:
        logger.info(message)


# (заголовок, anchor_key, depends_on, offset_days)
OPENING_TASK_TEMPLATE = [
    ("Учётки", None, None, None),
    ("Графики", None, None, None),
    ("Открытие", "opening", None, None),
    ("СКС", None, "vpk2", -3),
    ("Основной канал", None, None, None),
    ("Резервный канал", None, "vpk2", -5),
    ("Отправка оборудования", None, "vpk2", -5),
    ("Монтаж тумб", None, "vpk2", -3),
    ("КСО", None, None, None),
    ("ПНР", None, "vpk2", -2),
    ("ВПК-2", "vpk2", None, None),
    ("Пинпады", None, "vpk2", 1),
    ("ПНР КСО", None, "opening", -3),
    ("Проверка перед открытием", None, "opening", -2),
]

# (раздел, текст пункта) — шаблон чек-листа для "Открытия"
OPENING_CHECKLIST_TEMPLATE = [
    ("ПГЗ", "Редактирование файла ПГЗ"),
    ("Рассылка графиков", "КА"),
    ("Рассылка графиков", "Банк-эквайер"),
    ("Рассылка графиков", "Интернет — провайдер 1 и провайдер 2"),
    ("Рассылка графиков", "Оборудование"),
    ("Рассылка графиков", "АСЦН"),
    ("Учётки", "СБП"),
    ("Учётки", "Данные кристалл"),
    ("Учётки", "1771"),
    ("Учётки", "Редмул"),
    ("Учётки", "Плеер"),
    ("Учётки", "АДМ"),
    ("Учётки", "ОТРС"),
    ("Учётки", "Планировка"),
    ("Учётки", "Конфиги"),
    ("Учётки", "КПП"),
    ("Учётки", "Центрум"),
    ("Работы", "Прибытие оборудования"),
    ("Работы", "Монтаж каналов связи"),
    ("Работы", "Монтаж тумб"),
    ("Работы", "Монтаж СКС"),
    ("Работы", "ПНР / регистрация ключа / проверка работы ПК"),
    ("Работы", "Постановка на сервис новых магазинов"),
    ("Работы", "ВПК-2"),
    ("Работы", "Доставка терминалов"),
    ("ПНР КСО", "ПНР КСО"),
    ("ПНР КСО", "Проверка всех систем"),
    ("ПНР КСО", "Проверка обновлений касс"),
    ("ПНР КСО", "Проверка ОФД"),
    ("ПНР КСО", "Открытие"),
    ("ПГЗ", "Проверка файла ПГЗ"),
]

# (заголовок,) — обязательные задачи закрытия, без авторасчёта дат
CLOSING_TASK_TEMPLATE = [
    "Снятие с сервиса Кристалл",
    "Демонтаж банка-эквайера",
    "Демонтаж частичный КА",
    "Демонтаж полный КА",
    "Демонтаж провайдера 2",
    "Интернет демонтаж провайдера 1",
]

# (заголовок,) — обязательные задачи реконструкции вида "Полная", без авторасчёта дат
RECONSTRUCTION_FULL_TASK_TEMPLATE = [
    "Демонтаж частичный",
    "Демонтаж основной",
    "Рассылка графика работ",
    "Заказ оборудования",
    "Заказ пинпадов",
    "Письмо для восстановления канала",
    "Отправка оборудования на магазин",
    "Написать письмо для отправки",
    "СКС",
    "ПНР",
    "ВПК-2",
    "ПНР КСО",
    "Выезд для проверки",
]

# (заголовок,) — обязательные задачи доп. работ вида "Установка доп. КСО", без авторасчёта дат
EXTRA_KSO_TASK_TEMPLATE = ["СКС", "ПНР", "Доставка Пинпадов банка-эквайера"]


def recalc_dependents(store, anchor_key, anchor_date):
    """Пересчитывает даты всех задач магазина, привязанных к этой опорной дате."""
    if anchor_date is None:
        return
    dependents = Task.objects.filter(store=store, depends_on=anchor_key)
    for t in dependents:
        t.due_date = anchor_date + datetime.timedelta(days=t.offset_days or 0)
        t.save(update_fields=["due_date"])


def global_search(request):
    q = request.GET.get("q", "").strip()
    if len(q) < 2:
        return JsonResponse({"stores": [], "articles": [], "tasks": [], "notes": []})

    stores = Store.objects.filter(
        Q(number__icontains=q) | Q(address__icontains=q) | Q(region__icontains=q)
    )[:5]
    stores_result = [
        {
            "label": f"№{s.number} — {s.address}",
            "sublabel": s.get_branch_display() + (" | АРХИВ" if s.is_archived else ""),
            "url": reverse("store_detail", args=[s.id]),
        }
        for s in stores
    ]

    articles = Article.objects.filter(Q(title__icontains=q) | Q(content__icontains=q))[:5]
    articles_result = [
        {
            "label": a.title,
            "sublabel": "База знаний",
            "url": reverse("kb_detail", args=[a.id]),
        }
        for a in articles
    ]

    tasks = Task.objects.filter(title__icontains=q).select_related("store")[:5]
    tasks_result = [
        {
            "label": t.title,
            "sublabel": f"№{t.store.number} ({t.store.get_branch_display()})" + (" | АРХИВ" if t.store.is_archived else ""),
            "url": reverse("store_detail", args=[t.store_id]) + f"#task-{t.id}",
        }
        for t in tasks
    ]

    notes = Note.objects.filter(text__icontains=q, is_archived=False).select_related("task__store")[:5]
    notes_result = []
    for note in notes:
        label = " ".join(note.text.split())
        if len(label) > 60:
            label = label[:60] + "…"
        if note.task is None:
            url = reverse("notes") + f"#note-text-{note.id}"
            sublabel = "Заметка"
        else:
            url = reverse("store_detail", args=[note.task.store_id]) + f"#note-{note.id}"
            sublabel = f"№{note.task.store.number}" + (" | АРХИВ" if note.task.store.is_archived else "")
        notes_result.append({"label": label, "sublabel": sublabel, "url": url})

    return JsonResponse({
        "stores": stores_result,
        "articles": articles_result,
        "tasks": tasks_result,
        "notes": notes_result,
    })


def notifications_check(request):
    unread_count = Notification.objects.filter(is_read=False).count()
    items = [
        {
            "id": n.id,
            "text": n.text,
            "created_at": n.created_at.isoformat(),
        }
        for n in Notification.objects.all()[:20]
    ]
    return JsonResponse({"unread_count": unread_count, "items": items})


@require_POST
def notifications_mark_read(request):
    Notification.objects.filter(is_read=False).update(is_read=True)
    return JsonResponse({"ok": True})


def dashboard(request):
    today = datetime.date.today()
    tomorrow = today + datetime.timedelta(days=1)

    # Заметки без привязки к магазину (task=null) не привязаны ни к одной карточке
    # канбана — остаются отдельным блоком под канбаном, как и раньше.
    notes_overdue = Note.objects.filter(due_date__lt=today, is_archived=False, task__isnull=True).select_related("task__store")
    notes_today = Note.objects.filter(due_date=today, is_archived=False, task__isnull=True).select_related("task__store")
    notes_tomorrow = Note.objects.filter(due_date=tomorrow, is_archived=False, task__isnull=True).select_related("task__store")

    global_notes = Note.objects.filter(task__isnull=True, is_archived=False).prefetch_related("files")[:8]

    stores = Store.objects.filter(is_archived=False).prefetch_related("tasks")

    columns = {key: [] for key, _ in STORE_KANBAN_STATUS_CHOICES}
    kpi_overdue = kpi_today = kpi_tomorrow = 0
    for store in stores:
        tasks = list(store.tasks.all())
        store.tasks_overdue_count = sum(1 for t in tasks if t.due_date and t.due_date < today and t.status != "done")
        store.tasks_today_count = sum(1 for t in tasks if t.due_date == today and t.status != "done")
        store.tasks_tomorrow_count = sum(1 for t in tasks if t.due_date == tomorrow and t.status != "done")
        kpi_overdue += store.tasks_overdue_count
        kpi_today += store.tasks_today_count
        kpi_tomorrow += store.tasks_tomorrow_count
        columns.setdefault(store.kanban_status, []).append(store)

    # Прогресс за текущую неделю (пн-вс) — конвейерное производство,
    # общая готовность по всем задачам не показательна, а недельный план/факт да.
    week_start = today - datetime.timedelta(days=today.weekday())
    week_end = week_start + datetime.timedelta(days=6)
    tasks_this_week = (
        Task.objects.filter(store__is_archived=False, due_date__gte=week_start, due_date__lte=week_end)
        .select_related("store")
        .order_by("due_date", "order", "id")
    )
    kpi_week_total = tasks_this_week.count()
    kpi_week_done = sum(1 for t in tasks_this_week if t.status == "done")
    kpi_week_percent = round(kpi_week_done * 100 / kpi_week_total) if kpi_week_total else 0

    tasks_done_today = (
        Task.objects.filter(
            store__is_archived=False,
            status="done",
            completed_at__date=today,
        )
        .select_related("store")
        .order_by("-completed_at")
    )
    kpi_done_today = tasks_done_today.count()

    # Отдельный канбан задач по срочности (вкладка "Задачи" на главной), не связан с канбаном магазинов выше.
    active_tasks = (
        Task.objects.filter(store__is_archived=False)
        .exclude(status="done")
        .select_related("store")
        .order_by("due_date", "order", "id")
    )
    task_columns = {"overdue": [], "today": [], "tomorrow": [], "nodate": []}
    for t in active_tasks:
        if t.due_date is None:
            task_columns["nodate"].append(t)
        elif t.due_date < today:
            task_columns["overdue"].append(t)
        elif t.due_date == today:
            task_columns["today"].append(t)
        elif t.due_date == tomorrow:
            task_columns["tomorrow"].append(t)
        # остальные будущие даты (после завтра) в этот канбан не попадают

    # Календарь на месяц (виджет на дашборде). По умолчанию текущий месяц,
    # но принимает ?year=&month= для навигации вперёд/назад без AJAX —
    # объём данных маленький, поэтому просто перерисовываем страницу.
    try:
        cal_year = int(request.GET.get("year", today.year))
        cal_month = int(request.GET.get("month", today.month))
        # запас в год с каждого края, чтобы вычисление соседнего месяца
        # (prev/next) не вылетало за datetime.MINYEAR/MAXYEAR
        if not (datetime.MINYEAR + 1 <= cal_year <= datetime.MAXYEAR - 1):
            raise ValueError("год вне допустимого диапазона")
        datetime.date(cal_year, cal_month, 1)  # валидация диапазона
    except (TypeError, ValueError):
        cal_year, cal_month = today.year, today.month

    days_in_month = calendar.monthrange(cal_year, cal_month)[1]
    month_start = datetime.date(cal_year, cal_month, 1)
    month_end = datetime.date(cal_year, cal_month, days_in_month)

    cal_tasks = (
        Task.objects.filter(store__is_archived=False, due_date__range=(month_start, month_end))
        .select_related("store")
        .order_by("due_date", "due_time")
    )
    cal_notes = (
        Note.objects.filter(due_date__range=(month_start, month_end), is_archived=False)
        .exclude(task__store__is_archived=True)
        .select_related("task__store")
        .order_by("due_date", "due_time")
    )

    days_by_number = {day: {"tasks": [], "notes": []} for day in range(1, days_in_month + 1)}
    for t in cal_tasks:
        days_by_number[t.due_date.day]["tasks"].append(t)
    for n in cal_notes:
        days_by_number[n.due_date.day]["notes"].append(n)

    # Понедельник — первый день недели. weekday(): пн=0..вс=6.
    leading_blanks = month_start.weekday()
    calendar_days = [None] * leading_blanks
    for day in range(1, days_in_month + 1):
        date_obj = datetime.date(cal_year, cal_month, day)
        calendar_days.append({
            "day": day,
            "date": date_obj,
            "is_today": date_obj == today,
            "is_weekend": date_obj.weekday() >= 5,
            "tasks": days_by_number[day]["tasks"],
            "notes": days_by_number[day]["notes"],
            "has_items": bool(days_by_number[day]["tasks"] or days_by_number[day]["notes"]),
            "items_count": len(days_by_number[day]["tasks"]) + len(days_by_number[day]["notes"]),
        })
    # Дополняем сетку до полных недель (кратно 7), чтобы ровно рисовать строки.
    while len(calendar_days) % 7:
        calendar_days.append(None)
    calendar_weeks = [calendar_days[i:i + 7] for i in range(0, len(calendar_days), 7)]

    prev_month_date = month_start - datetime.timedelta(days=1)
    next_month_date = month_end + datetime.timedelta(days=1)

    return render(request, "board/dashboard.html", {
        "columns": columns,
        "kpi_stores_count": len(stores),
        "kpi_overdue": kpi_overdue,
        "kpi_today": kpi_today,
        "kpi_tomorrow": kpi_tomorrow,
        "kpi_week_percent": kpi_week_percent,
        "tasks_this_week": tasks_this_week,
        "kpi_done_today": kpi_done_today,
        "tasks_done_today": tasks_done_today,
        "notes_overdue": notes_overdue,
        "notes_today": notes_today,
        "notes_tomorrow": notes_tomorrow,
        "global_notes": global_notes,
        "branches": BRANCH_CHOICES,
        "kanban_statuses": STORE_KANBAN_STATUS_CHOICES,
        "task_columns": task_columns,
        "calendar_weeks": calendar_weeks,
        "calendar_year": cal_year,
        "calendar_month": cal_month,
        "calendar_month_name": MONTH_NAMES_RU[cal_month - 1],
        "calendar_prev_year": prev_month_date.year,
        "calendar_prev_month": prev_month_date.month,
        "calendar_next_year": next_month_date.year,
        "calendar_next_month": next_month_date.month,
        "calendar_is_current_month": (cal_year, cal_month) == (today.year, today.month),
        "calendar_keep_open": bool(request.GET.get("year") or request.GET.get("month")),
    })


def weekly_summary(request):
    today = datetime.date.today()
    week_start = today - datetime.timedelta(days=today.weekday() + 7)
    week_end = week_start + datetime.timedelta(days=6)

    # anchor_key="opening" — общий "ключ" итоговой задачи магазина: у типа
    # "Открытие" его проставляет шаблон OPENING_TASK_TEMPLATE, у остальных
    # типов его вручную ставит пользователь кнопкой "Сделать ключевой"
    # (toggle_task_anchor) на нужной задаче.
    tasks = (
        Task.objects.filter(anchor_key="opening", due_date__gte=week_start, due_date__lte=week_end)
        .select_related("store")
        .prefetch_related("store__problems")
        .order_by("store__store_type", "due_date", "store__number")
    )

    columns = {key: [] for key, _ in STORE_TYPE_CHOICES}
    for t in tasks:
        columns[t.store.store_type].append(t)

    return render(request, "board/weekly_summary.html", {
        "columns": [(key, label, columns[key]) for key, label in STORE_TYPE_CHOICES],
        "week_start": week_start,
        "week_end": week_end,
    })


def _autolink_anchor_by_title(task):
    # Автопривязка "ключевой" задачи по слову "открытие" в названии (любой
    # регистр) — чтобы не заставлять пользователя каждый раз жать "Сделать
    # ключевой" руками. Работает для любого типа магазина: у магазинов типа
    # "Новые Магазины", созданных через автошаблон, ключ уже стоит на задаче
    # "Открытие" (тогда ничего не делаем — см. проверку ниже), но не все
    # магазины заводятся через шаблон, поэтому исключения по типу магазина
    # быть не должно. Не перебиваем уже выбранный ключ.
    if "открытие" not in task.title.lower():
        return
    if Task.objects.filter(store=task.store, anchor_key="opening").exclude(id=task.id).exists():
        return
    task.anchor_key = "opening"
    task.save(update_fields=["anchor_key"])


def toggle_task_anchor(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    if request.method == "POST":
        if task.anchor_key == "opening":
            task.anchor_key = None
            task.save(update_fields=["anchor_key"])
        else:
            Task.objects.filter(store=task.store).exclude(id=task.id).update(anchor_key=None)
            task.anchor_key = "opening"
            task.save(update_fields=["anchor_key"])
    return redirect("store_detail", store_id=task.store_id)


def branch_view(request, branch, store_type):
    if request.method == "POST":
        number = request.POST.get("number", "").strip()
        if number:
            store = Store.objects.create(
                branch=branch,
                store_type=store_type,
                number=number,
                address=request.POST.get("address", "").strip(),
                region=request.POST.get("region", "").strip(),
                contact=request.POST.get("contact", "").strip(),
            )
            if store_type == "opening":
                Task.objects.bulk_create([
                    Task(
                        store=store,
                        title=title,
                        anchor_key=anchor_key,
                        depends_on=depends_on,
                        offset_days=offset_days,
                        order=i,
                    )
                    for i, (title, anchor_key, depends_on, offset_days) in enumerate(OPENING_TASK_TEMPLATE)
                ])

                store.contractors = request.POST.getlist("contractors")
                store.kso_count = request.POST.get("kso_count", "").strip()
                store.cash_count = request.POST.get("cash_count", "").strip()
                store.area_format = request.POST.get("area_format", "").strip()
                store.columns_count = request.POST.get("columns_count", "").strip()
                store.columns_type = request.POST.get("columns_type", "").strip()
                store.has_meat_scale = request.POST.get("has_meat_scale") == "on"
                store.access_points_count = request.POST.get("access_points_count", "").strip()
                store.license_number = request.POST.get("license_number", "").strip()
                store.kpp = request.POST.get("kpp", "").strip()
                store.cluster = request.POST.get("cluster", "").strip()
                store.ibp_count = request.POST.get("ibp_count", "").strip()
                store.provider = request.POST.get("provider", "").strip()
                store.save()
                # router_values считаем до записи строки в письма, чтобы столбец
                # "Внешние IP резервного канала" (autofill=wan_ip2_reserve) заполнился сразу.
                router_values = compute_router_config_for_store(store)
                # Строку в письмах создаём до простановки опорных дат, чтобы дозаполнение
                # дат ниже нашло куда писать (иначе строка ещё не существует).
                sync_new_store_to_letters(store, router_values)

                if router_values:
                    StoreRouterConfig.objects.update_or_create(store=store, defaults={"values": router_values})
                    device_values = compute_device_ips_for_network(router_values.get("NETWORK"))
                    if device_values:
                        StoreDeviceIPConfig.objects.update_or_create(store=store, defaults={"values": device_values})

                try:
                    kso_total = int(store.kso_count)
                except (TypeError, ValueError):
                    kso_total = 0
                if kso_total > 0:
                    StoreCashRegister.objects.bulk_create([
                        StoreCashRegister(store=store, order=i, number=str(KSO_NUMBER_BASE + i))
                        for i in range(kso_total)
                    ])

                for anchor_key, field_name in [("opening", "opening_date"), ("vpk2", "vpk2_date")]:
                    date_str = request.POST.get(field_name, "").strip()
                    if not date_str:
                        continue
                    try:
                        parsed_date = datetime.date.fromisoformat(date_str)
                    except ValueError:
                        continue
                    anchor_task = Task.objects.filter(store=store, anchor_key=anchor_key).first()
                    if not anchor_task:
                        continue
                    anchor_task.due_date = parsed_date
                    anchor_task.save(update_fields=["due_date"])
                    recalc_dependents(store, anchor_key, parsed_date)
            elif store_type == "closing":
                Task.objects.bulk_create([
                    Task(store=store, title=title, order=i)
                    for i, title in enumerate(CLOSING_TASK_TEMPLATE)
                ])
            elif store_type == "reconstruction":
                store.reconstruction_kind = request.POST.get("reconstruction_kind", "").strip()
                store.save(update_fields=["reconstruction_kind"])
                if store.reconstruction_kind == "full":
                    Task.objects.bulk_create([
                        Task(store=store, title=title, order=i)
                        for i, title in enumerate(RECONSTRUCTION_FULL_TASK_TEMPLATE)
                    ])
            elif store_type == "extra":
                store.extra_kind = request.POST.get("extra_kind", "").strip()
                store.save(update_fields=["extra_kind"])
                if store.extra_kind == "extra_kso":
                    Task.objects.bulk_create([
                        Task(store=store, title=title, order=i)
                        for i, title in enumerate(EXTRA_KSO_TASK_TEMPLATE)
                    ])
        return redirect("branch", branch=branch, store_type=store_type)

    show_archived = request.GET.get("archived") == "1"
    stores = Store.objects.filter(branch=branch, store_type=store_type, is_archived=show_archived)
    branch_label = dict(BRANCH_CHOICES).get(branch, branch)
    type_label = dict(STORE_TYPE_CHOICES).get(store_type, store_type)

    return render(request, "board/branch.html", {
        "stores": stores,
        "show_archived": show_archived,
        "branch": branch,
        "store_type": store_type,
        "branch_label": branch_label,
        "type_label": type_label,
        "branches": BRANCH_CHOICES,
        "letter_contractors": LetterCompany.objects.filter(is_store_contractor=True) if store_type == "opening" else None,
        "area_format_choices": AREA_FORMAT_CHOICES,
        "contact_choices": CONTRACTOR_CHOICES,
    })


def _excel_date(value):
    return value.strftime("%d.%m.%Y") if value else ""


def _excel_task_due(tasks_by_title, title):
    task = tasks_by_title.get(title)
    return _excel_date(task.due_date) if task else ""


def _lan_mask_bits(mask_str):
    """255.255.255.128 -> 25. Возвращает None, если маска пустая/не парсится."""
    if not mask_str:
        return None
    try:
        return ipaddress.IPv4Network(f"0.0.0.0/{mask_str}").prefixlen
    except ValueError:
        return None


def _cell(value):
    # Заменяем переносы строк и табуляции на пробел — иначе многострочные
    # поля (например "Примечания") при вставке TSV в Excel сдвигают
    # колонки или создают лишнюю строку.
    return " ".join(str(value or "").split())


def _build_store_excel_row(store):
    """Собирает строку из 39 значений для вставки в рабочий Excel-файл пользователя
    (кнопка "Скопировать строку для Excel" на карточке магазина) — порядок колонок
    A-AM жёстко фиксирован под шаблон пользователя, см. план функции."""
    top_tasks = store.tasks.filter(parent__isnull=True)
    tasks_by_title = {t.title: t for t in top_tasks}
    tasks_by_anchor = {t.anchor_key: t for t in top_tasks if t.anchor_key}

    router_config = StoreRouterConfig.objects.filter(store=store).first()
    values = router_config.values if router_config else {}

    network = values.get("NETWORK", "")
    lan_mask_bits = _lan_mask_bits(values.get("LANMASK"))
    if network and lan_mask_bits is not None:
        internal_ips = f"{network}/{lan_mask_bits}"
    else:
        internal_ips = network

    wanip2 = values.get("WANIP2", "")
    wanmask2 = values.get("WANMASK2", "")
    wangw2 = values.get("WANGW2", "")
    if wanip2:
        parts = [wanip2]
        if wanmask2:
            parts[0] = f"{wanip2}/{wanmask2}"
        if wangw2:
            external_ips2 = f"{parts[0]}, шлюз {wangw2}"
        else:
            external_ips2 = parts[0]
    else:
        external_ips2 = ""

    opening_task = tasks_by_anchor.get("opening")
    vpk2_task = tasks_by_anchor.get("vpk2")

    return [_cell(v) for v in [
        "",  # A №п/п
        store.get_branch_display(),  # B Филиал
        store.number,  # C № объекта
        store.license_number,  # D Лицензия для касс
        store.kpp,  # E КПП
        store.address,  # F Адрес
        "",  # G Дата выхода строителей
        store.region,  # H Регион
        store.cluster,  # I Кластер
        store.cash_count,  # J Кол-во касс
        store.kso_count,  # K Кол-во КСО
        store.ibp_count,  # L Кол-во ИБП
        store.access_points_count,  # M Кол-во ТД WiFi
        "да" if store.has_meat_scale else "нет",  # N Кол-во весов с чекопечатью
        store.area_format,  # O Формат по площади
        store.columns_count,  # P Кол-во колонок
        store.columns_type,  # Q Тип колонок
        values.get("TUNIP", ""),  # R TUNIP
        internal_ips,  # S Внутренние IP адреса
        _excel_task_due(tasks_by_title, "Основной канал"),  # T Дата монтажа канала связи
        values.get("WANIP", ""),  # U Внешние IP адреса
        store.provider,  # V Провайдер
        values.get("TUNIP2", ""),  # W Туннельные IP резервного канала
        _excel_task_due(tasks_by_title, "Резервный канал"),  # X Дата монтажа резервного канала
        external_ips2,  # Y Внешние IP резервного канала
        _excel_task_due(tasks_by_title, "Отправка оборудования"),  # Z Дата отправки оборудования
        _excel_task_due(tasks_by_title, "СКС"),  # AA Дата СКС
        _excel_task_due(tasks_by_title, "ПНР"),  # AB Дата ПНР
        _excel_task_due(tasks_by_title, "Монтаж тумб"),  # AC Монтаж тумб для КСО
        _excel_task_due(tasks_by_title, "ПНР КСО"),  # AD ПНР КСО
        _excel_task_due(tasks_by_title, "Пинпады"),  # AE Дата монтажа пин-падов
        _excel_date(vpk2_task.due_date) if vpk2_task else "",  # AF Дата ВПК-2
        _excel_task_due(tasks_by_title, "Проверка перед открытием"),  # AG Дата проверки по чек-листам
        _excel_date(opening_task.due_date) if opening_task else "",  # AH Дата открытия
        store.notes,  # AI Примечания
        store.contact,  # AJ Подрядчик СКС
        store.contact,  # AK Подрядчик ПНР
        store.contact,  # AL Подрядчик КТО
        "",  # AM Дата закрытия
    ]]


def _build_ascn_rows(store):
    """Собирает строки (по одной на каждую кассу магазина) из 19 значений A-S
    для вставки в таблицу АСЦН (кнопка "Скопировать строки для АСЦН" на
    карточке магазина)."""
    router_config = StoreRouterConfig.objects.filter(store=store).first()
    network = router_config.values.get("NETWORK", "") if router_config else ""

    device_config = StoreDeviceIPConfig.objects.filter(store=store).first()
    device_values = device_config.values if device_config else {}
    dev_pc1 = device_values.get("DEV_PC1", "")
    dev_pc2 = device_values.get("DEV_PC2", "")
    device_ip = device_values.get("DEV_SERVER_KSO", "")

    utm_addr = f"http://{dev_pc1}:8080/xml" if dev_pc1 else ""
    setmark_addr = f"http://{dev_pc2}:9000" if dev_pc2 else ""
    lmch_addr = f"{dev_pc1}:5995" if dev_pc1 else ""
    mask_dns = f"maska {DEVICE_MASK} DNS1 {DEVICE_DNS1} DNS 2 {DEVICE_DNS2}"

    rows = []
    for reg in store.cash_registers.all():
        kso_ips = compute_kso_ips(network, reg.order)
        if kso_ips:
            kso_ips_str = f"IP СБ {kso_ips['ip_sb']}  IP ФР {kso_ips['ip_fr']} GW {kso_ips['gw']}"
        else:
            kso_ips_str = ""
        rows.append([_cell(v) for v in [
            store.number,  # A № объекта
            store.get_branch_display(),  # B Филиал
            store.region,  # C Регион
            store.locality,  # D Населённый пункт
            store.address,  # E Адрес
            device_ip,  # F IP модуля интеграции
            COMPANY_INN,  # G ИНН
            store.kpp,  # H КПП
            utm_addr,  # I Адрес УТМ
            setmark_addr,  # J Адрес setMark
            lmch_addr,  # K ЛМЧЗ
            reg.number,  # L Номер КСО
            f"КСО {reg.number}" if reg.number else "",  # M Наименование КСО
            kso_ips_str,  # N IP СБ/ФР/GW
            mask_dns,  # O Маска и DNS
            reg.loymax_data,  # P LOYMAX posid/login/password
            store.fsrar_id,  # Q FSRAR_ID
            reg.sbp_terminal,  # R Номер терминала СБП
            reg.sbp_link,  # S Кассовая ссылка СБП
        ]])
    return rows


def store_detail(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    tasks = (
        store.tasks.filter(parent__isnull=True)
        .prefetch_related("notes__files", "subtasks__notes__files")
        .order_by("order", "id")
    )
    report = store.reports.prefetch_related("photos").order_by("-created_at").first()
    contacts = store.contacts.all()
    checklist_items = list(store.checklist_items.all())
    checklist_categories = []
    for item in checklist_items:
        if item.category not in checklist_categories:
            checklist_categories.append(item.category)
    checklist_items.sort(key=lambda item: (checklist_categories.index(item.category), item.order, item.id))
    logs = store.logs.prefetch_related("files")
    problems = store.problems.all()
    router_placeholders, _ = _router_config_placeholders(store)
    router_config = StoreRouterConfig.objects.filter(store=store).first()
    has_network = bool(router_config and router_config.values.get("NETWORK"))
    device_config = StoreDeviceIPConfig.objects.filter(store=store).first()
    device_values = device_config.values if device_config else {}
    device_ip_fields = [
        {"key": key, "label": label, "value": device_values.get(key, "")}
        for key, label in DEVICE_IP_LABELS.items()
    ]
    excel_row_tsv = "\t".join(_build_store_excel_row(store))

    ascn_registers = list(store.cash_registers.all())
    ascn_network = router_config.values.get("NETWORK", "") if router_config else ""
    dev_pc1 = device_values.get("DEV_PC1", "")
    dev_pc2 = device_values.get("DEV_PC2", "")
    ascn_device_ip = device_values.get("DEV_SERVER_KSO", "")
    ascn_utm_addr = f"http://{dev_pc1}:8080/xml" if dev_pc1 else ""
    ascn_setmark_addr = f"http://{dev_pc2}:9000" if dev_pc2 else ""
    ascn_lmch_addr = f"{dev_pc1}:5995" if dev_pc1 else ""
    ascn_rows_display = []
    for reg in ascn_registers:
        kso_ips = compute_kso_ips(ascn_network, reg.order)
        ascn_rows_display.append({
            "obj": reg,
            "kso_name": f"КСО {reg.number}" if reg.number else "",
            "ip_sb": kso_ips["ip_sb"] if kso_ips else "",
            "ip_fr": kso_ips["ip_fr"] if kso_ips else "",
            "gw": kso_ips["gw"] if kso_ips else "",
        })
    ascn_rows_tsv = "\n".join("\t".join(row) for row in _build_ascn_rows(store))

    return render(request, "board/store_detail.html", {
        "store": store,
        "tasks": tasks,
        "report": report,
        "contacts": contacts,
        "checklist_items": checklist_items,
        "logs": logs,
        "problems": problems,
        "branches": BRANCH_CHOICES,
        "router_placeholders": router_placeholders,
        "device_ip_fields": device_ip_fields,
        "has_network": has_network,
        "excel_row_tsv": excel_row_tsv,
        "ascn_rows_display": ascn_rows_display,
        "ascn_rows_tsv": ascn_rows_tsv,
        "company_inn": COMPANY_INN,
        "ascn_utm_addr": ascn_utm_addr,
        "ascn_setmark_addr": ascn_setmark_addr,
        "ascn_lmch_addr": ascn_lmch_addr,
        "ascn_device_ip": ascn_device_ip,
    })


def archive_store(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST":
        store.is_archived = not store.is_archived
        store.save(update_fields=["is_archived"])
        if store.is_archived:
            return redirect("branch", branch=store.branch, store_type=store.store_type)
    return redirect("store_detail", store_id=store.id)


def duplicate_store(request, store_id):
    original = get_object_or_404(Store, id=store_id)
    if request.method != "POST":
        return redirect("store_detail", store_id=original.id)

    new_store = Store.objects.create(
        branch=original.branch,
        store_type=original.store_type,
        number=original.number + " (копия)",
        address=original.address,
        region=original.region,
        contact=original.contact,
        notes=original.notes,
    )
    id_map = {}
    for task in original.tasks.filter(parent__isnull=True):
        new_task = Task.objects.create(
            store=new_store,
            title=task.title,
            due_date=task.due_date,
            order=task.order,
            anchor_key=task.anchor_key,
            depends_on=task.depends_on,
            offset_days=task.offset_days,
        )
        id_map[task.id] = new_task.id
    for task in original.tasks.filter(parent__isnull=False):
        Task.objects.create(
            store=new_store,
            parent_id=id_map.get(task.parent_id),
            title=task.title,
            due_date=task.due_date,
            order=task.order,
        )
    return redirect("store_detail", store_id=new_store.id)


def add_task(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST":
        title = request.POST.get("title", "").strip()
        due_date = request.POST.get("due_date") or None
        due_time = request.POST.get("due_time") or None
        icon_name = request.POST.get("icon_name", "").strip()
        icon_color = request.POST.get("icon_color", "").strip()
        if icon_name not in VALID_ICONS or icon_color not in VALID_COLORS:
            icon_name, icon_color = "", ""
        if title:
            last_order = store.tasks.filter(parent__isnull=True).count()
            task = Task.objects.create(
                store=store, title=title, due_date=due_date, due_time=due_time, order=last_order,
                icon_name=icon_name, icon_color=icon_color,
            )
            _autolink_anchor_by_title(task)
    return redirect("store_detail", store_id=store.id)


def update_task_icon(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    if request.method == "POST":
        icon_name = request.POST.get("icon_name", "").strip()
        icon_color = request.POST.get("icon_color", "").strip()
        if icon_name not in VALID_ICONS or icon_color not in VALID_COLORS:
            icon_name, icon_color = "", ""  # "" + "" = вернуться к автоподбору
        task.icon_name = icon_name
        task.icon_color = icon_color
        task.save(update_fields=["icon_name", "icon_color"])
    return redirect("store_detail", store_id=task.store_id)


def add_subtask(request, task_id):
    parent = get_object_or_404(Task, id=task_id)
    if request.method == "POST":
        title = request.POST.get("title", "").strip()
        if title:
            task = Task.objects.create(
                store=parent.store,
                parent=parent,
                title=title,
                order=parent.subtasks.count(),
            )
            _autolink_anchor_by_title(task)
    return redirect("store_detail", store_id=parent.store_id)


def update_task_status(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    if request.method == "POST":
        status = request.POST.get("status")
        valid_statuses = {c[0] for c in Task._meta.get_field("status").choices}
        if status in valid_statuses:
            task.status = status
            task.save()
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        store = task.store
        today = datetime.date.today()
        tomorrow = today + datetime.timedelta(days=1)
        tasks = list(store.tasks.all())
        return JsonResponse({
            "ok": True,
            "task_id": task.id,
            "task_status": task.status,
            "store_id": store.id,
            "progress_percent": store.progress_percent,
            "tasks_overdue_count": sum(1 for t in tasks if t.due_date and t.due_date < today and t.status != "done"),
            "tasks_today_count": sum(1 for t in tasks if t.due_date == today and t.status != "done"),
            "tasks_tomorrow_count": sum(1 for t in tasks if t.due_date == tomorrow and t.status != "done"),
        })
    if request.POST.get("next") == "dashboard":
        return redirect("dashboard")
    return redirect("store_detail", store_id=task.store_id)


def store_tasks_panel(request, store_id):
    """Список задач магазина для боковой панели на канбане главной страницы (без перехода на страницу магазина)."""
    store = get_object_or_404(Store, id=store_id)
    tasks = store.tasks.filter(parent__isnull=True).order_by("order", "id")
    return render(request, "board/partials/store_tasks_panel.html", {
        "store": store,
        "tasks": tasks,
    })


@require_POST
def update_store_kanban_status(request, store_id):
    """Drag-and-drop карточки магазина между колонками канбана на главной странице."""
    store = get_object_or_404(Store, id=store_id)
    status = request.POST.get("status")
    valid_statuses = {c[0] for c in Store._meta.get_field("kanban_status").choices}
    if status not in valid_statuses:
        return JsonResponse({"ok": False, "error": "Некорректная стадия"}, status=400)
    store.kanban_status = status
    store.save(update_fields=["kanban_status"])
    return JsonResponse({"ok": True})


def update_task_due_date(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    if request.method == "POST":
        due_date = request.POST.get("due_date") or None
        has_time_field = "due_time" in request.POST
        due_time = request.POST.get("due_time") or None
        old_date = task.due_date.isoformat() if task.due_date else None
        old_time = task.due_time.strftime("%H:%M") if task.due_time else None
        changed = (due_date != old_date) or (has_time_field and due_time != old_time)
        task.due_date = due_date
        update_fields = ["due_date"]
        if has_time_field:
            task.due_time = due_time
            update_fields.append("due_time")
        if changed:
            task.notified_at = None  # срок поменялся — уведомление о новом времени можно слать заново
            update_fields.append("notified_at")
        task.save(update_fields=update_fields)
        if task.anchor_key and due_date:
            parsed = datetime.date.fromisoformat(due_date)
            recalc_dependents(task.store, task.anchor_key, parsed)
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        store = task.store
        today = datetime.date.today()
        tomorrow = today + datetime.timedelta(days=1)
        tasks = list(store.tasks.all())
        return JsonResponse({
            "ok": True,
            "task_id": task.id,
            "due_date": task.due_date.isoformat() if task.due_date else None,
            "store_id": store.id,
            "progress_percent": store.progress_percent,
            "tasks_overdue_count": sum(1 for t in tasks if t.due_date and t.due_date < today and t.status != "done"),
            "tasks_today_count": sum(1 for t in tasks if t.due_date == today and t.status != "done"),
            "tasks_tomorrow_count": sum(1 for t in tasks if t.due_date == tomorrow and t.status != "done"),
        })
    return redirect("store_detail", store_id=task.store_id)


def move_task(request, task_id, direction):
    task = get_object_or_404(Task, id=task_id)
    siblings = list(
        Task.objects.filter(store=task.store, parent=task.parent).order_by("order", "id")
    )
    idx = siblings.index(task)
    swap_idx = idx - 1 if direction == "up" else idx + 1
    if request.method == "POST" and 0 <= swap_idx < len(siblings):
        other = siblings[swap_idx]
        task.order, other.order = other.order, task.order
        task.save(update_fields=["order"])
        other.save(update_fields=["order"])
    return redirect("store_detail", store_id=task.store_id)


def delete_task(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    store_id = task.store_id
    if request.method == "POST":
        task.delete()
    return redirect("store_detail", store_id=store_id)


# ------------------------- ЖУРНАЛ ПО ОБЪЕКТУ -------------------------

def add_store_log(request, store_id):
    """Новая запись журнала (с сайта)."""
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST":
        text = request.POST.get("text", "").strip()
        if text:
            log = StoreLog.objects.create(store=store, text=text, source="web")
            for f in request.FILES.getlist("files"):
                StoreLogFile.objects.create(log=log, file=f)
    return redirect("store_detail", store_id=store.id)


def edit_store_log(request, log_id):
    log = get_object_or_404(StoreLog, id=log_id)
    if request.method == "POST":
        text = request.POST.get("text", "").strip()
        if text:
            log.text = text
            log.save(update_fields=["text"])
    return redirect("store_detail", store_id=log.store_id)


def delete_store_log(request, log_id):
    log = get_object_or_404(StoreLog, id=log_id)
    store_id = log.store_id
    if request.method == "POST":
        log.delete()
    return redirect("store_detail", store_id=store_id)


def add_store_log_file(request, log_id):
    log = get_object_or_404(StoreLog, id=log_id)
    if request.method == "POST":
        for f in request.FILES.getlist("files"):
            StoreLogFile.objects.create(log=log, file=f)
    return redirect("store_detail", store_id=log.store_id)


def delete_store_log_file(request, file_id):
    lf = get_object_or_404(StoreLogFile, id=file_id)
    store_id = lf.log.store_id
    if request.method == "POST":
        lf.delete()
    return redirect("store_detail", store_id=store_id)


# ------------------------- ОТЧЁТ (ОДИН НА МАГАЗИН) -------------------------

def _format_logs_for_report(store):
    """Склеивает записи журнала в читаемый текст, от старых к новым."""
    entries = store.logs.order_by("created_at")
    lines = []
    for lg in entries:
        lines.append(f"{lg.created_at:%d.%m.%Y %H:%M}\n{lg.text}\n")
    return "\n".join(lines).strip()


def collect_report(request, store_id):
    """Создаёт отчёт (если его ещё нет) или пересобирает текст из журнала заново."""
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST":
        report = store.reports.order_by("-created_at").first()
        collected = _format_logs_for_report(store)
        if report is None:
            report = Report.objects.create(store=store, final_text=collected)
        else:
            report.final_text = collected
            report.save(update_fields=["final_text"])
    return redirect("store_detail", store_id=store.id)


def save_report(request, report_id):
    """Сохраняет текст отчёта после ручной правки."""
    report = get_object_or_404(Report, id=report_id)
    if request.method == "POST":
        report.final_text = request.POST.get("final_text", "").strip()
        report.save(update_fields=["final_text"])
    return redirect("store_detail", store_id=report.store_id)


def add_photo(request, report_id):
    report = get_object_or_404(Report, id=report_id)
    if request.method == "POST":
        for f in request.FILES.getlist("photos"):
            ReportPhoto.objects.create(report=report, file=f)
    return redirect("store_detail", store_id=report.store_id)


def delete_photo(request, photo_id):
    photo = get_object_or_404(ReportPhoto, id=photo_id)
    store_id = photo.report.store_id
    if request.method == "POST":
        photo.delete()
    return redirect("store_detail", store_id=store_id)


def letters(request):
    companies = list(LetterCompany.objects.all())
    initial_state = {
        "companyOrder": [c.key for c in companies],
        "companies": {
            c.key: {
                "name": c.name,
                "page": c.page,
                "header": c.header,
                "subject": c.subject,
                "recipients": c.recipients,
                "columns": c.columns,
                "rows": c.rows,
                "isStoreContractor": c.is_store_contractor,
            }
            for c in companies
        },
    }
    opening_task_titles = [title for title, *_ in OPENING_TASK_TEMPLATE]
    return render(request, "board/letters.html", {
        "initial_state": initial_state,
        "opening_task_titles": opening_task_titles,
    })


@require_POST
def letters_save(request):
    try:
        payload = json.loads(request.body.decode("utf-8"))
    except (ValueError, UnicodeDecodeError):
        return JsonResponse({"ok": False, "error": "bad json"}, status=400)

    company_order = payload.get("companyOrder") or []
    companies = payload.get("companies") or {}

    existing_keys = set(LetterCompany.objects.values_list("key", flat=True))
    payload_keys = set(companies.keys())
    LetterCompany.objects.filter(key__in=existing_keys - payload_keys).delete()

    for index, key in enumerate(company_order):
        data = companies.get(key)
        if not data:
            continue
        page = data.get("page") or "schedules"
        if page not in ("schedules", "accounts"):
            page = "schedules"
        LetterCompany.objects.update_or_create(
            key=key,
            defaults={
                "name": data.get("name") or key,
                "page": page,
                "order": index,
                "header": data.get("header") or "",
                "subject": data.get("subject") or "",
                "recipients": data.get("recipients") or "",
                "columns": data.get("columns") or [],
                "rows": data.get("rows") or [],
                "is_store_contractor": bool(data.get("isStoreContractor")),
            },
        )
    return JsonResponse({"ok": True})


def notes_page(request):
    notes = Note.objects.filter(task__isnull=True, is_archived=False).prefetch_related("files")
    return render(request, "board/notes.html", {"notes": notes, "branches": BRANCH_CHOICES})


def add_note(request):
    if request.method == "POST":
        text = request.POST.get("text", "").strip()
        if text:
            note = Note.objects.create(
                text=text,
                due_date=request.POST.get("due_date") or None,
                due_time=request.POST.get("due_time") or None,
            )
            for f in request.FILES.getlist("files"):
                NoteFile.objects.create(note=note, file=f)
    return redirect("notes")


def add_task_note(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    if request.method == "POST":
        text = request.POST.get("text", "").strip()
        if text:
            note = Note.objects.create(
                text=text,
                task=task,
                due_date=request.POST.get("due_date") or None,
                due_time=request.POST.get("due_time") or None,
            )
            for f in request.FILES.getlist("files"):
                NoteFile.objects.create(note=note, file=f)
    return redirect("store_detail", store_id=task.store_id)


def edit_note(request, note_id):
    note = get_object_or_404(Note, id=note_id)
    if request.method == "POST":
        text = request.POST.get("text", "").strip()
        if text:
            note.text = text
        new_due_date = request.POST.get("due_date") or None
        new_due_time = request.POST.get("due_time") or None
        if new_due_date != (note.due_date.isoformat() if note.due_date else None) or \
           new_due_time != (note.due_time.isoformat(timespec="minutes") if note.due_time else None):
            note.notified_at = None
        note.due_date = new_due_date
        note.due_time = new_due_time
        note.save()
    if request.POST.get("next") == "dashboard":
        return redirect("dashboard")
    if note.task_id:
        return redirect("store_detail", store_id=note.task.store_id)
    return redirect("notes")


def delete_note(request, note_id):
    note = get_object_or_404(Note, id=note_id)
    store_id = note.task.store_id if note.task_id else None
    if request.method == "POST":
        note.delete()
    if request.POST.get("next") == "dashboard":
        return redirect("dashboard")
    if store_id:
        return redirect("store_detail", store_id=store_id)
    return redirect("notes")


def add_note_file(request, note_id):
    note = get_object_or_404(Note, id=note_id)
    if request.method == "POST":
        for f in request.FILES.getlist("files"):
            NoteFile.objects.create(note=note, file=f)
    if note.task_id:
        return redirect("store_detail", store_id=note.task.store_id)
    return redirect("notes")


def delete_note_file(request, file_id):
    nf = get_object_or_404(NoteFile, id=file_id)
    note = nf.note
    if request.method == "POST":
        nf.delete()
    if note.task_id:
        return redirect("store_detail", store_id=note.task.store_id)
    return redirect("notes")


def contacts_page(request):
    contacts = Contact.objects.select_related("store").all()
    stores = Store.objects.filter(is_archived=False).order_by("branch", "number")

    grouped = {code: [] for code, _label in BRANCH_CHOICES}
    no_branch = []
    for contact in contacts:
        branch = contact.effective_branch
        if branch in grouped:
            grouped[branch].append(contact)
        else:
            no_branch.append(contact)

    groups = [
        {"code": code, "label": label, "contacts": grouped[code]}
        for code, label in BRANCH_CHOICES
    ]
    if no_branch:
        groups.append({"code": "", "label": "Без филиала", "contacts": no_branch})

    return render(
        request,
        "board/contacts.html",
        {
            "contacts": contacts,
            "branches": BRANCH_CHOICES,
            "stores": stores,
            "groups": groups,
            "active_branch": request.GET.get("branch", ""),
        },
    )


def add_contact(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST":
        name = request.POST.get("name", "").strip()
        if name:
            Contact.objects.create(
                store=store,
                name=name,
                role=request.POST.get("role", "").strip(),
                phone=request.POST.get("phone", "").strip(),
                telegram=request.POST.get("telegram", "").strip(),
                max_contact=request.POST.get("max_contact", "").strip(),
            )
    next_url = request.POST.get("next", "")
    if next_url == "contacts":
        return redirect("contacts")
    return redirect("store_detail", store_id=store.id)


def add_contact_no_store(request):
    if request.method == "POST":
        name = request.POST.get("name", "").strip()
        branch = request.POST.get("branch", "").strip()
        if name and branch:
            Contact.objects.create(
                branch=branch,
                name=name,
                role=request.POST.get("role", "").strip(),
                phone=request.POST.get("phone", "").strip(),
                telegram=request.POST.get("telegram", "").strip(),
                max_contact=request.POST.get("max_contact", "").strip(),
            )
    return redirect("contacts")


def delete_contact(request, contact_id):
    contact = get_object_or_404(Contact, id=contact_id)
    store_id = contact.store_id
    if request.method == "POST":
        contact.delete()
    next_url = request.POST.get("next", "")
    if store_id is None:
        return redirect("contacts")
    if next_url == "contacts":
        return redirect("contacts")
    return redirect("store_detail", store_id=store_id)


def add_problem(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST":
        text = request.POST.get("text", "").strip()
        if text:
            Problem.objects.create(store=store, text=text)
    return redirect("store_detail", store_id=store.id)


def toggle_problem(request, problem_id):
    problem = get_object_or_404(Problem, id=problem_id)
    if request.method == "POST":
        problem.is_resolved = not problem.is_resolved
        problem.resolved_at = timezone.now() if problem.is_resolved else None
        problem.save(update_fields=["is_resolved", "resolved_at"])
    return redirect("store_detail", store_id=problem.store_id)


def delete_problem(request, problem_id):
    problem = get_object_or_404(Problem, id=problem_id)
    store_id = problem.store_id
    if request.method == "POST":
        problem.delete()
    return redirect("store_detail", store_id=store_id)


def _match_branch(question):
    """Если в вопросе явно назван филиал (код или название) — вернуть его код."""
    ql = question.lower()
    for code, label in BRANCH_CHOICES:
        if code in ql.split() or label.lower() in ql:
            return code
    return None


def _match_store_ids(question, stores):
    """Найти магазины, чей номер объекта явно упомянут в вопросе (для подробностей)."""
    ql = question.lower()
    matched = []
    for store in stores:
        number = (store.number or "").strip().lower()
        if number and len(number) >= 2 and number in ql:
            matched.append(store.id)
    return matched


def _match_date_range(question):
    """Если в вопросе явно названы даты (диапазон вида 27.07.26-02.08.26 или "с X по Y")
    — вернуть (начало, конец). Локальная модель ненадёжно ищет нужные строки в длинном
    плоском списке задач "на глаз" — с датами код фильтрует их точно сам, а модели
    останется только оформить готовый список в ответ."""
    matches = re.findall(r"(\d{1,2})\.(\d{1,2})\.(\d{2,4})", question)
    dates = []
    for d, m, y in matches:
        y = int(y)
        if y < 100:
            y += 2000
        try:
            dates.append(datetime.date(y, int(m), int(d)))
        except ValueError:
            continue
    if not dates:
        return None
    return min(dates), max(dates)


def _match_relative_week(question):
    """Если явно спрашивают про "прошлую неделю" — вернуть (понедельник, воскресенье)
    той календарной недели, без необходимости печатать даты руками. Триггер для
    готового еженедельного отчёта в _build_ai_context (см. ниже)."""
    ql = question.lower()
    if "прошл" not in ql or "недел" not in ql:
        return None
    today = timezone.localdate()
    start_this_week = today - datetime.timedelta(days=today.weekday())
    start = start_this_week - datetime.timedelta(days=7)
    end = start_this_week - datetime.timedelta(days=1)
    return start, end


def _store_compact_line(store):
    """Одна строка на магазин: без задач/журнала/заметок целиком — экономит токены,
    но даёт статус, счётчики и ближайший срок, чего достаточно для сводных вопросов
    ("статусы для начальников")."""
    tasks = list(store.tasks.all())
    total = len(tasks)
    done = sum(1 for t in tasks if t.status == "done")
    waiting = sum(1 for t in tasks if t.status == "waiting")
    progress = sum(1 for t in tasks if t.status == "progress")
    open_dates = sorted(t.due_date for t in tasks if t.status != "done" and t.due_date)
    next_due = open_dates[0].strftime("%d.%m.%Y") if open_dates else "-"
    archived_label = " | АРХИВ" if store.is_archived else ""

    # Опорные задачи (anchor_key) — это конкретно "Открытие" и "ВПК-2", их даты часто
    # спрашивают напрямую по названию/диапазону дат, а не через "ближайший срок"
    # (задача может быть уже отмечена "готово" и не попасть в next_due).
    anchors = {t.anchor_key: t for t in tasks if t.anchor_key}
    anchor_bits = []
    for key, label in ANCHOR_CHOICES:
        t = anchors.get(key)
        due = t.due_date.strftime("%d.%m.%Y") if t and t.due_date else "-"
        anchor_bits.append(f"{label}: {due}")

    return (
        f"МАГАЗИН {store.number} | {store.get_branch_display()} | {store.get_store_type_display()}"
        f"{archived_label} | задач: {total} (готово {done}, в работе {progress}, ждём {waiting}) | "
        f"ближайший срок: {next_due} | " + " | ".join(anchor_bits)
    )


def _store_detail_lines(store):
    """Полная информация по одному магазину: задачи, подзадачи, заметки, журнал, контакты —
    только для магазинов, явно упомянутых в вопросе, чтобы не тащить это по всей базе."""
    lines = []
    archived_label = " | АРХИВ (закрыт/завершён)" if store.is_archived else ""
    lines.append(
        f"ПОДРОБНО: МАГАЗИН {store.number} | {store.get_branch_display()} | "
        f"{store.get_store_type_display()} | адрес: {store.address or '-'} | "
        f"регион: {store.region or '-'} | контакт: {store.contact or '-'}{archived_label}"
    )
    for c in store.contacts.all():
        lines.append(f"  доп.контакт: {c.name} ({c.role}) {c.phone} {c.telegram}".rstrip())
    for t in store.tasks.filter(parent__isnull=True):
        due = t.due_date.strftime("%d.%m.%Y") if t.due_date else "-"
        lines.append(f"  задача: {t.title} | статус: {t.get_status_display()} | срок: {due}")
        for n in t.notes.filter(is_archived=False):
            lines.append(f"    заметка к задаче «{t.title}»: {n.text[:300]}")
        for st in t.subtasks.all():
            sdue = st.due_date.strftime("%d.%m.%Y") if st.due_date else "-"
            lines.append(f"    подзадача: {st.title} | статус: {st.get_status_display()} | срок: {sdue}")
            for n in st.notes.filter(is_archived=False):
                lines.append(f"      заметка к подзадаче «{st.title}»: {n.text[:300]}")
    store_logs = list(store.logs.all()[:40])
    if store_logs:
        lines.append("  ЖУРНАЛ РАБОТ (свежие записи сверху):")
        for lg in store_logs:
            when = lg.created_at.strftime("%d.%m.%Y %H:%M")
            lines.append(f"    [{when}] {lg.text[:500]}")
    return lines


def _build_ai_context(question="", max_chars=150000):
    """Сводка по всем магазинам (компактно, по одной строке) + подробности только по тем
    магазинам, что явно упомянуты в вопросе по номеру, и/или отфильтрованные по филиалу,
    если он назван. Так модель видит всю базу целиком (ничего не обрезается по объёму),
    а токены тратятся только на детали там, где они реально нужны."""
    stores_qs = Store.objects.prefetch_related(
        "tasks", "tasks__notes", "tasks__subtasks", "tasks__subtasks__notes", "contacts", "logs"
    )
    branch = _match_branch(question) if question else None
    if branch:
        stores_qs = stores_qs.filter(branch=branch)
    stores = list(stores_qs)

    lines = ["СВОДКА ПО ВСЕМ МАГАЗИНАМ:"]
    for store in stores:
        lines.append(_store_compact_line(store))

    # Если в вопросе явно названы даты — отфильтровать задачи по сроку кодом (точно,
    # без ошибок) и отдать модели готовый список первым делом. Локальная модель на
    # длинном плоском списке ниже периодически "не находит" реально существующие
    # строки — это подстраховка от таких промахов.
    date_range = _match_date_range(question) if question else None
    if date_range:
        start, end = date_range
        lines.append("")
        lines.append(
            f"ТОЧНЫЙ СПИСОК ЗАДАЧ СО СРОКОМ ИМЕННО В ПЕРИОДЕ {start.strftime('%d.%m.%Y')}"
            f"–{end.strftime('%d.%m.%Y')} (отфильтровано кодом по датам, а не моделью — "
            "для ответа на вопрос про период используй ИМЕННО этот список, не ищи заново "
            "по общему списку ниже):"
        )
        found_any = False
        for store in stores:
            for t in store.tasks.all():
                if t.due_date and start <= t.due_date <= end:
                    lines.append(
                        f"  МАГАЗИН {store.number} | {t.title} | {t.get_status_display()} | "
                        f"срок: {t.due_date.strftime('%d.%m.%Y')}"
                    )
                    found_any = True
        if not found_any:
            lines.append("  (в этом периоде задач с указанным сроком нет)")

    # Ключевой запрос "отчёт за прошлую неделю" — та же идея, что и с точными датами выше:
    # период считается кодом (не моделью), задачи и проблемы отфильтрованы точно, включая
    # архивные магазины и все типы работ (открытие/реконструкция/доп.работы/закрытие).
    week_range = _match_relative_week(question) if question else None
    if week_range:
        w_start, w_end = week_range
        lines.append("")
        lines.append(
            f"ЕЖЕНЕДЕЛЬНЫЙ ОТЧЁТ ЗА ПРОШЛУЮ НЕДЕЛЮ ({w_start.strftime('%d.%m.%Y')}"
            f"–{w_end.strftime('%d.%m.%Y')}, отфильтровано кодом — используй именно эти "
            "данные для ответа, включая архивные магазины и все типы работ):"
        )
        lines.append("  ЗАДАЧИ СО СРОКОМ НА ЭТОЙ НЕДЕЛЕ:")
        found_any = False
        for store in stores:
            for t in store.tasks.all():
                if t.due_date and w_start <= t.due_date <= w_end:
                    archived_bit = " | АРХИВ" if store.is_archived else ""
                    lines.append(
                        f"    МАГАЗИН {store.number} | {store.get_branch_display()} | "
                        f"{store.get_store_type_display()}{archived_bit} | {t.title} | "
                        f"{t.get_status_display()} | срок: {t.due_date.strftime('%d.%m.%Y')}"
                    )
                    found_any = True
        if not found_any:
            lines.append("    (задач со сроком на этой неделе нет)")

        new_problems_qs = Problem.objects.filter(
            created_at__date__gte=w_start, created_at__date__lte=w_end
        ).select_related("store")
        open_problems_qs = Problem.objects.filter(is_resolved=False).select_related("store")
        if branch:
            new_problems_qs = new_problems_qs.filter(store__branch=branch)
            open_problems_qs = open_problems_qs.filter(store__branch=branch)

        lines.append("  НОВЫЕ ПРОБЛЕМЫ, ЗАВЕДЁННЫЕ ЗА ЭТУ НЕДЕЛЮ:")
        new_problems = list(new_problems_qs)
        if new_problems:
            for p in new_problems:
                status = "решено" if p.is_resolved else "НЕ решено"
                lines.append(f"    МАГАЗИН {p.store.number} | {p.text} | {status}")
        else:
            lines.append("    (новых проблем за эту неделю не заводили)")

        lines.append("  ВСЕ НЕРЕШЁННЫЕ ПРОБЛЕМЫ НА СЕГОДНЯ (не только за эту неделю):")
        open_problems = list(open_problems_qs)
        if open_problems:
            for p in open_problems:
                lines.append(
                    f"    МАГАЗИН {p.store.number} | {p.text} | заведена {p.created_at.strftime('%d.%m.%Y')}"
                )
        else:
            lines.append("    (нерешённых проблем нет)")

    # Плоский список ВСЕХ задач (включая подзадачи) по всем магазинам — название, статус,
    # срок, без заметок/журнала/контактов. Нужен для вопросов вида "какие задачи со статусом
    # готово были с 27.07 по 02.08" — сводки по счётчикам для такого недостаточно, а полный
    # дамп с журналом (как раньше) был слишком тяжёлым по токенам.
    lines.append("")
    lines.append("ВСЕ ЗАДАЧИ (магазин | название | статус | срок):")
    for store in stores:
        for t in store.tasks.all():
            due = t.due_date.strftime("%d.%m.%Y") if t.due_date else "-"
            lines.append(f"  {store.number} | {t.title} | {t.get_status_display()} | {due}")

    matched_ids = set(_match_store_ids(question, stores)) if question else set()
    if matched_ids:
        lines.append("")
        for store in stores:
            if store.id in matched_ids:
                lines.extend(_store_detail_lines(store))

    lines.append("")
    lines.append("ОБЩИЕ ЗАМЕТКИ:")
    for n in Note.objects.filter(task__isnull=True, is_archived=False)[:50]:
        lines.append(f"  - {n.text[:200]}")

    text = "\n".join(lines)
    truncated = False
    if len(text) > max_chars:
        text = text[:max_chars]
        truncated = True
    return text, truncated


class _FakeResponse:
    """Оборачивает ответ Ollama в интерфейс (.status_code/.json()) с той же
    формой JSON (candidates/content/parts), которую ожидает код ниже по
    стеку (ai_ask/ai_test/kb_ai_search)."""

    def __init__(self, status_code, payload):
        self.status_code = status_code
        self._payload = payload

    def json(self):
        return self._payload


def _call_ollama(prompt, timeout=300):
    is_thinking_model = "qwen3" in settings.OLLAMA_MODEL.lower()
    payload = {
        "model": settings.OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        # num_ctx=6144 — на видеокарте 8ГБ (RTX 3070) модель 9b сама весит ~7.2ГБ,
        # с num_ctx=12288 места под KV-кэш не хватало и часть слоёв уходила на CPU
        # (Ollama показывала 28%/72% CPU/GPU) — отсюда таймауты по 300 сек. При 6144
        # модель влезает в видеопамять целиком.
        "options": {"num_ctx": 6144, "num_predict": 2048, "temperature": 0.2},
    }
    if is_thinking_model:
        # Qwen3 — "думающая" модель: даём ей честно порассуждать (think:true) перед
        # ответом, это заметно улучшает качество на вопросах вроде "что сделано за
        # период". За это платим временем — таймаут увеличен отдельно ниже.
        payload["think"] = True
    try:
        r = requests.post(
            f"{settings.OLLAMA_URL}/api/generate",
            json=payload,
            timeout=timeout,
        )
    except requests.RequestException as e:
        return _FakeResponse(502, {"error": {"message": f"Ollama недоступна: {e}"}})

    if r.status_code != 200:
        return _FakeResponse(r.status_code, {"error": {"message": r.text[:300]}})

    answer = r.json().get("response", "").strip()

    if not answer and is_thinking_model:
        # Размышления съели весь num_predict, ответа не осталось — вместо пустоты
        # пробуем ещё раз без "размышлений", отдав весь бюджет токенов сразу под ответ.
        logger.warning("Qwen3 отдал пустой ответ (думал до конца лимита токенов), пробую без think")
        payload_retry = dict(payload)
        payload_retry["think"] = False
        try:
            r2 = requests.post(
                f"{settings.OLLAMA_URL}/api/generate",
                json=payload_retry,
                timeout=timeout,
            )
            if r2.status_code == 200:
                answer = r2.json().get("response", "").strip()
        except requests.RequestException:
            pass

    return _FakeResponse(200, {"candidates": [{"content": {"parts": [{"text": answer}]}}]})


def _call_llm(prompt, timeout=30):
    # "Думающей" модели (Qwen3) на размышления перед ответом нужно больше времени —
    # даём ей до 10 минут, пользователь согласен подождать ради более полного ответа.
    ollama_timeout = 600 if "qwen3" in settings.OLLAMA_MODEL.lower() else 300
    return _call_ollama(prompt, timeout=max(timeout, ollama_timeout))


def ai_ask(request):
    if request.method != "POST":
        return JsonResponse({"error": "method not allowed"}, status=405)

    log = []
    if not _llm_ready():
        _log(log, _llm_not_ready_message())
        return JsonResponse({"error": _llm_not_ready_message(), "log": log}, status=400)

    question = request.POST.get("question", "").strip()
    if not question:
        return JsonResponse({"error": "Пустой вопрос", "log": log}, status=400)

    history_raw = request.POST.get("history", "")
    history_lines = []
    if history_raw:
        try:
            history_items = json.loads(history_raw)
            for item in history_items:
                role = "Пользователь" if item.get("role") == "user" else "Ассистент"
                text = str(item.get("text", "")).strip()
                if text:
                    history_lines.append(f"{role}: {text}")
        except (ValueError, TypeError):
            pass

    logger.info("=== Новый вопрос к базе: %s", question)
    _log(log, "Собираю данные по базе (магазины, задачи, заметки)...")
    # лимит контекста Ollama ограничен видеопамятью/num_ctx локальной модели
    max_chars = 12000
    context_text, truncated = _build_ai_context(question, max_chars=max_chars)
    _log(log, f"Собрано {len(context_text)} символов контекста" + (" (обрезано по лимиту)" if truncated else ""))

    history_block = ""
    if history_lines:
        history_block = "ПРЕДЫДУЩИЙ ДИАЛОГ (для контекста):\n" + "\n".join(history_lines) + "\n\n"

    prompt = (
        "Ты — помощник по учёту магазинов, задач и заметок. Ниже — текущие данные из базы. "
        "Отвечай по делу, основываясь только на этих данных, ничего не придумывай. "
        "Если под вопрос подходит несколько магазинов/задач — перечисли их все, по одному "
        "пункту на строку, и для каждого обязательно укажи номер магазина, суть задачи и "
        "дату/срок — не пропускай дату, даже если пунктов много. "
        "Отвечай ТОЛЬКО на русском языке. Ни при каких обстоятельствах не используй "
        "китайский, английский или любой другой язык. "
        "Если данных не хватает для ответа — так и скажи (на русском).\n\n"
        f"ДАННЫЕ:\n{context_text}\n\n{history_block}ВОПРОС: {question}"
    )

    _log(log, f"Отправляю запрос модели ({_active_llm_name()})...")
    try:
        resp = _call_llm(prompt)
    except requests.RequestException as e:
        _log(log, f"Сетевая ошибка: {e}")
        return JsonResponse({"error": "Не удалось связаться с ИИ", "log": log}, status=502)

    _log(log, f"Ответ получен, HTTP {resp.status_code}")
    try:
        data = resp.json()
    except ValueError:
        _log(log, "Не удалось разобрать ответ (не JSON)")
        return JsonResponse({"error": "Некорректный ответ от ИИ", "log": log}, status=502)

    if resp.status_code != 200:
        err = data.get("error", {}).get("message", "неизвестная ошибка")
        _log(log, f"Ошибка API: {err}")
        return JsonResponse({"error": err, "log": log}, status=502)

    candidates = data.get("candidates", [])
    if not candidates:
        _log(log, "Пустой ответ (нет candidates) — возможно, сработал фильтр безопасности")
        return JsonResponse({"error": "Пустой ответ от ИИ", "log": log}, status=502)

    parts = candidates[0].get("content", {}).get("parts", [])
    answer = "".join(p.get("text", "") for p in parts).strip()
    _log(log, "Готово")
    logger.info("ОТВЕТ: %s", answer)
    return JsonResponse({"answer": answer, "log": log})


def ai_test(request):
    if not _llm_ready():
        return JsonResponse({"ok": False, "error": _llm_not_ready_message()})
    try:
        resp = _call_llm("Ответь одним словом: работает", timeout=20)
        data = resp.json()
    except requests.RequestException as e:
        return JsonResponse({"ok": False, "error": f"Сетевая ошибка: {e}"})
    except ValueError:
        return JsonResponse({"ok": False, "error": "Некорректный ответ от ИИ"})

    if resp.status_code != 200:
        return JsonResponse({"ok": False, "error": data.get("error", {}).get("message", str(resp.status_code))})
    return JsonResponse({"ok": True, "model": _active_llm_name()})


def kb_list(request):
    query = request.GET.get("q", "").strip()
    context = {"query": query}
    if query:
        context["articles"] = Article.objects.filter(title__icontains=query)
    else:
        context["categories"] = KbCategory.objects.all()
        context["has_uncategorized"] = Article.objects.filter(category__isnull=True).exists()
    return render(request, "board/kb_list.html", context)


def kb_category_detail(request, category_id):
    if category_id == 0:
        category = None
        articles = Article.objects.filter(category__isnull=True)
    else:
        category = get_object_or_404(KbCategory, id=category_id)
        articles = Article.objects.filter(category=category)
    category_files = KbCategoryFile.objects.filter(category=category)
    return render(request, "board/kb_category_detail.html", {
        "category": category,
        "articles": articles,
        "category_files": category_files,
        "category_id_for_url": category.id if category else 0,
    })


def kb_category_add(request):
    if request.method == "POST":
        name = request.POST.get("name", "").strip()
        image = request.FILES.get("image")
        if not name or not image:
            return render(request, "board/kb_category_form.html", {"error": "Укажите название и картинку", "name": name})
        KbCategory.objects.create(name=name, image=image, image_has_caption=False)
        return redirect("kb_list")
    return render(request, "board/kb_category_form.html")


def kb_detail(request, article_id):
    article = get_object_or_404(Article, id=article_id)
    return render(request, "board/kb_detail.html", {"article": article})


def kb_add(request):
    category_id = _clean_category_id(request.GET.get("category_id"))
    draft = Article.objects.create(title="Новая статья", content="", category_id=category_id)
    return redirect("kb_edit", article_id=draft.id)


def _run_to_html(run):
    from django.utils.html import escape
    text = escape(run.text)
    if not text:
        return ""
    if run.bold:
        text = f"<b>{text}</b>"
    if run.italic:
        text = f"<i>{text}</i>"
    if run.underline:
        text = f"<u>{text}</u>"
    return text


def _extract_run_images(run, docx_document, article):
    from docx.oxml.ns import qn
    tags = []
    blips = run._element.findall(".//" + qn("a:blip"))
    for blip in blips:
        rId = blip.get(qn("r:embed"))
        if not rId:
            continue
        try:
            image_part = docx_document.part.related_parts[rId]
        except KeyError:
            continue
        from django.core.files.base import ContentFile
        img_obj = ArticleImage(article=article)
        fname = rId + ".png"
        img_obj.image.save(fname, ContentFile(image_part.blob), save=True)
        tags.append(f'<img src="{img_obj.image.url}" style="max-width:100%;display:block;margin:10px 0;">')
    return tags


def _paragraph_to_html(paragraph, docx_document, article):
    parts = []
    for run in paragraph.runs:
        parts.append(_run_to_html(run))
        parts.extend(_extract_run_images(run, docx_document, article))
    return "".join(parts)


def _docx_to_html(file_obj, article):
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = DocxDocument(file_obj)
    html_parts = []
    title_guess = None
    list_mode = None  # 'ul' / 'ol' / None — чтобы соседние пункты списка объединялись в один список

    def close_list():
        nonlocal list_mode
        if list_mode:
            html_parts.append(f"</{list_mode}>")
            list_mode = None

    def iter_block_items(document):
        body = document.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, document)
            elif child.tag == qn("w:tbl"):
                yield Table(child, document)

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            style = (block.style.name or "").lower()
            inner_html = _paragraph_to_html(block, doc, article)
            has_text = bool(block.text.strip())
            has_image = "<img" in inner_html
            if not has_text and not has_image:
                continue

            if "heading 1" in style or style == "title":
                close_list()
                html_parts.append(f"<h2>{inner_html}</h2>")
                if title_guess is None:
                    title_guess = block.text.strip()
            elif "heading 2" in style:
                close_list()
                html_parts.append(f"<h3>{inner_html}</h3>")
            elif "heading" in style:
                close_list()
                html_parts.append(f"<h4>{inner_html}</h4>")
            elif "list bullet" in style:
                if list_mode != "ul":
                    close_list()
                    html_parts.append("<ul>")
                    list_mode = "ul"
                html_parts.append(f"<li>{inner_html}</li>")
            elif "list number" in style:
                if list_mode != "ol":
                    close_list()
                    html_parts.append("<ol>")
                    list_mode = "ol"
                html_parts.append(f"<li>{inner_html}</li>")
            else:
                close_list()
                html_parts.append(f"<p>{inner_html}</p>")
        else:
            close_list()
            rows_html = []
            for row in block.rows:
                cells_html = "".join(
                    f'<td style="border:1px solid #999;padding:6px;">{cell.text}</td>' for cell in row.cells
                )
                rows_html.append(f"<tr>{cells_html}</tr>")
            html_parts.append('<table style="border-collapse:collapse;width:100%;">' + "".join(rows_html) + "</table>")

    close_list()
    return title_guess, "\n".join(html_parts)


def _pdf_to_html(file_obj):
    from pypdf import PdfReader
    reader = PdfReader(file_obj)
    title_guess = None
    html_parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if title_guess is None:
                title_guess = line
            from django.utils.html import escape
            html_parts.append(f"<p>{escape(line)}</p>")
    return title_guess, "\n".join(html_parts)


def _clean_category_id(raw):
    try:
        category_id = int(raw)
    except (TypeError, ValueError):
        return None
    if not KbCategory.objects.filter(id=category_id).exists():
        return None
    return category_id


def kb_import(request):
    if request.method == "POST":
        category_id = _clean_category_id(request.POST.get("category_id"))

        f = request.FILES.get("file")
        if not f:
            return render(request, "board/kb_import.html", {"error": "Выбери файл", "category_id": category_id})

        name = f.name.lower()
        if not (name.endswith(".docx") or name.endswith(".pdf")):
            return render(request, "board/kb_import.html", {"error": "Поддерживаются только .docx и .pdf (старый .doc не подходит)", "category_id": category_id})

        article = Article.objects.create(title=f.name.rsplit(".", 1)[0], content="", category_id=category_id)
        try:
            if name.endswith(".docx"):
                title_guess, html = _docx_to_html(f, article)
            else:
                title_guess, html = _pdf_to_html(f)
        except Exception as e:
            article.delete()
            logger.exception("Ошибка импорта файла в базу знаний")
            return render(request, "board/kb_import.html", {"error": f"Не удалось разобрать файл: {e}", "category_id": category_id})

        if title_guess:
            article.title = title_guess[:250]
        article.content = html
        article.save()
        return redirect("kb_edit", article_id=article.id)

    category_id = _clean_category_id(request.GET.get("category_id"))
    return render(request, "board/kb_import.html", {"category_id": category_id})


def kb_edit(request, article_id):
    article = get_object_or_404(Article, id=article_id)
    categories = KbCategory.objects.all()
    if request.method == "POST":
        title = request.POST.get("title", "").strip()
        content = request.POST.get("content", "")
        if title:
            article.title = title
            article.content = content
            article.category_id = _clean_category_id(request.POST.get("category_id"))
            article.save()
            return redirect("kb_detail", article_id=article.id)
    return render(request, "board/kb_form.html", {"article": article, "categories": categories})


def kb_delete(request, article_id):
    article = get_object_or_404(Article, id=article_id)
    category_id = article.category_id or 0
    if request.method == "POST":
        article.delete()
    return redirect("kb_category_detail", category_id=category_id)


def kb_upload_image(request, article_id):
    article = get_object_or_404(Article, id=article_id)
    if request.method != "POST" or "image" not in request.FILES:
        return JsonResponse({"error": "Нет файла"}, status=400)
    img = ArticleImage.objects.create(article=article, image=request.FILES["image"])
    return JsonResponse({"url": img.image.url})


def kb_add_file(request, article_id):
    article = get_object_or_404(Article, id=article_id)
    if request.method == "POST":
        for f in request.FILES.getlist("files"):
            ArticleFile.objects.create(article=article, file=f)
    return redirect("kb_detail", article_id=article.id)


def kb_delete_file(request, file_id):
    f = get_object_or_404(ArticleFile, id=file_id)
    article_id = f.article_id
    if request.method == "POST":
        f.delete()
    return redirect("kb_detail", article_id=article_id)


def kb_category_add_file(request, category_id):
    category = get_object_or_404(KbCategory, id=category_id) if category_id != 0 else None
    if request.method == "POST":
        for f in request.FILES.getlist("files"):
            KbCategoryFile.objects.create(category=category, file=f)
    return redirect("kb_category_detail", category_id=category_id)


def kb_category_delete_file(request, file_id):
    f = get_object_or_404(KbCategoryFile, id=file_id)
    category_id = f.category_id or 0
    if request.method == "POST":
        f.delete()
    return redirect("kb_category_detail", category_id=category_id)


def _resolve_media_path(src):
    """По src картинки ("/media/..." или "media/...") возвращает путь к файлу на диске, если он существует."""
    if not src:
        return None
    media_url = settings.MEDIA_URL.lstrip("/")
    stripped_src = src.lstrip("/")
    if not stripped_src.startswith(media_url):
        return None
    local_path = settings.MEDIA_ROOT / stripped_src[len(media_url):]
    return local_path if local_path.exists() else None


def _add_runs_from_node(paragraph, node, bold=False, italic=False, underline=False):
    from bs4 import NavigableString, Tag
    from docx.shared import Inches
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip() or " " in text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
        elif isinstance(child, Tag):
            if child.name == "img":
                local_path = _resolve_media_path(child.get("src", ""))
                if local_path:
                    try:
                        paragraph.add_run().add_picture(str(local_path), width=Inches(5.5))
                    except Exception:
                        pass
                continue
            b = bold or child.name in ("b", "strong")
            i = italic or child.name in ("i", "em")
            u = underline or child.name == "u"
            if child.name == "br":
                paragraph.add_run().add_break()
            else:
                _add_runs_from_node(paragraph, child, b, i, u)


def _html_to_docx(html_content, title):
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(title, level=1)

    soup = BeautifulSoup(html_content or "", "html.parser")
    top_nodes = soup.contents

    def handle_node(node):
        from bs4 import Tag
        if not isinstance(node, Tag):
            return
        if node.name in ("h1", "h2", "h3"):
            level = int(node.name[1]) + 1
            doc.add_heading(node.get_text(), level=min(level, 4))
        elif node.name == "img":
            local_path = _resolve_media_path(node.get("src", ""))
            if local_path:
                try:
                    doc.add_picture(str(local_path), width=Inches(5.5))
                except Exception:
                    pass
        elif node.name in ("ul", "ol"):
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet" if node.name == "ul" else "List Number")
                _add_runs_from_node(p, li)
        elif node.name == "table":
            rows = node.find_all("tr")
            if rows:
                cols = max(len(r.find_all(["td", "th"])) for r in rows)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                for r in rows:
                    cells = r.find_all(["td", "th"])
                    row_cells = table.add_row().cells
                    for idx, c in enumerate(cells):
                        row_cells[idx].text = c.get_text()
        elif node.name in ("p", "div"):
            p = doc.add_paragraph()
            _add_runs_from_node(p, node)
        else:
            p = doc.add_paragraph()
            _add_runs_from_node(p, node)

    for node in top_nodes:
        handle_node(node)

    return doc


def kb_download_docx(request, article_id):
    article = get_object_or_404(Article, id=article_id)
    doc = _html_to_docx(article.content, article.title)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{article.title}.docx".replace("/", "-")
    return FileResponse(buf, as_attachment=True, filename=filename)


def kb_download_txt(request, article_id):
    from bs4 import BeautifulSoup
    article = get_object_or_404(Article, id=article_id)
    text = BeautifulSoup(article.content or "", "html.parser").get_text("\n")
    content = f"{article.title}\n{'=' * len(article.title)}\n\n{text}"
    response = HttpResponse(content, content_type="text/plain; charset=utf-8")
    filename = f"{article.title}.txt".replace("/", "-")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def kb_print(request, article_id):
    article = get_object_or_404(Article, id=article_id)
    return render(request, "board/kb_print.html", {"article": article})


def kb_ai_search(request):
    if request.method != "POST":
        return JsonResponse({"error": "method not allowed"}, status=405)

    log = []
    if not _llm_ready():
        _log(log, _llm_not_ready_message())
        return JsonResponse({"error": _llm_not_ready_message(), "log": log}, status=400)

    question = request.POST.get("question", "").strip()
    if not question:
        return JsonResponse({"error": "Пустой вопрос", "log": log}, status=400)

    logger.info("=== Поиск по базе знаний: %s", question)
    articles = list(Article.objects.all().values("id", "title", "content"))
    if not articles:
        _log(log, "В базе знаний пока нет статей")
        return JsonResponse({"error": "В базе знаний пока нет статей", "log": log}, status=400)

    _log(log, f"Просматриваю {len(articles)} статей базы знаний...")

    listing_parts = []
    for a in articles:
        import re
        plain = re.sub("<[^>]+>", " ", a["content"] or "")[:600]
        listing_parts.append(f'ID {a["id"]}: "{a["title"]}"\nТекст: {plain}')
    listing = "\n\n".join(listing_parts)

    prompt = (
        "Ниже — список статей базы знаний (ID, заголовок, отрывок текста). "
        "Пользователь ищет инструкцию по своему вопросу. "
        "Определи ОДНУ наиболее подходящую статью. "
        "Отвечай ТОЛЬКО на русском языке, не используй китайский или английский. "
        "Ответь строго в формате:\nID: <номер или 0 если ничего не подходит>\nПОЧЕМУ: <короткое объяснение на русском>\n\n"
        f"СТАТЬИ:\n{listing}\n\nВОПРОС: {question}"
    )

    _log(log, f"Отправляю запрос модели ({_active_llm_name()})...")
    try:
        resp = _call_llm(prompt)
    except requests.RequestException as e:
        _log(log, f"Сетевая ошибка: {e}")
        return JsonResponse({"error": "Не удалось связаться с ИИ", "log": log}, status=502)

    _log(log, f"Ответ получен, HTTP {resp.status_code}")
    try:
        data = resp.json()
    except ValueError:
        _log(log, "Не удалось разобрать ответ (не JSON)")
        return JsonResponse({"error": "Некорректный ответ от ИИ", "log": log}, status=502)

    if resp.status_code != 200:
        err = data.get("error", {}).get("message", "неизвестная ошибка")
        _log(log, f"Ошибка API: {err}")
        return JsonResponse({"error": err, "log": log}, status=502)

    candidates = data.get("candidates", [])
    if not candidates:
        _log(log, "Пустой ответ от ИИ")
        return JsonResponse({"error": "Пустой ответ от ИИ", "log": log}, status=502)

    parts = candidates[0].get("content", {}).get("parts", [])
    answer_text = "".join(p.get("text", "") for p in parts).strip()

    import re
    id_match = re.search(r"ID:\s*(\d+)", answer_text)
    reason_match = re.search(r"ПОЧЕМУ:\s*(.+)", answer_text, re.DOTALL)
    article_id = int(id_match.group(1)) if id_match else 0
    reason = reason_match.group(1).strip() if reason_match else answer_text

    _log(log, "Готово")

    if not article_id:
        logger.info("РЕЗУЛЬТАТ: статья не найдена. %s", reason)
        return JsonResponse({"found": False, "reason": reason, "log": log})

    article = Article.objects.filter(id=article_id).first()
    if not article:
        logger.info("РЕЗУЛЬТАТ: ИИ указал несуществующий ID %s", article_id)
        return JsonResponse({"found": False, "reason": reason, "log": log})

    logger.info("РЕЗУЛЬТАТ: статья «%s» (id=%s). %s", article.title, article.id, reason)
    return JsonResponse({"found": True, "article_id": article.id, "title": article.title, "reason": reason, "log": log})


# Сайт, Telegram-бот, голосовой ассистент и почтовый ассистент — независимые
# долгоживущие процессы, каждый пишет в свой файл лога (см. settings.py,
# почему они разведены). Здесь просто общий список для переключателя на
# странице /logs/ — ключ используется в URL (?file=...), реальное имя файла
# берётся строго из этого словаря, а не из пользовательского ввода.
LOG_FILES = {
    "ai": ("ai.log", "Сайт"),
    "bot": ("bot.log", "Telegram-бот"),
    "voice": ("voice_assistant.log", "Голосовой ассистент"),
    "email": ("email_watcher.log", "Почта"),
}


def logs_page(request):
    # Файлы переворачиваются каждую полночь (см. settings.py), так что обычно в них
    # и так только сегодняшние записи — но фильтруем по дате в начале строки на случай,
    # если что-то осталось от предыдущего дня (например, процесс не перезапускался).
    file_key = request.GET.get("file", "ai")
    if file_key not in LOG_FILES:
        file_key = "ai"
    filename, _ = LOG_FILES[file_key]
    log_path = settings.LOGS_DIR / filename
    today_prefix = datetime.date.today().strftime("%d.%m.%Y")
    lines = []
    if log_path.exists():
        with open(log_path, encoding="utf-8", errors="replace") as f:
            lines = [ln for ln in f if ln.startswith(today_prefix)]
    tabs = [{"key": key, "label": label} for key, (_, label) in LOG_FILES.items()]
    return render(request, "board/logs.html", {
        "log_text": "".join(lines),
        "tabs": tabs,
        "current_file": file_key,
    })


def logs_download(request):
    file_key = request.GET.get("file", "ai")
    if file_key not in LOG_FILES:
        file_key = "ai"
    filename, _ = LOG_FILES[file_key]
    log_path = settings.LOGS_DIR / filename
    if not log_path.exists():
        return HttpResponse("Лог пока пуст.", content_type="text/plain; charset=utf-8")
    return FileResponse(open(log_path, "rb"), as_attachment=True, filename=filename)


def _build_today_digest():
    today = datetime.date.today()
    tasks_today = Task.objects.filter(
        due_date=today, store__is_archived=False
    ).exclude(status="done").select_related("store")
    tasks_overdue = Task.objects.filter(
        due_date__lt=today, store__is_archived=False
    ).exclude(status="done").select_related("store")

    lines = [f"📋 Задачи на {today.strftime('%d.%m.%Y')}"]

    if tasks_today:
        lines.append("")
        for t in tasks_today:
            lines.append(f"• {t.title} — {t.store.get_branch_display()} · {t.store.number}")
    else:
        lines.append("\nНа сегодня открытых задач нет.")

    if tasks_overdue:
        lines.append("\n⚠️ Просрочено:")
        for t in tasks_overdue:
            lines.append(f"• {t.title} — {t.store.get_branch_display()} · {t.store.number} (срок {t.due_date.strftime('%d.%m.%Y')})")

    return "\n".join(lines)


def send_telegram_message(text):
    if not settings.TELEGRAM_BOT_TOKEN or not settings.TELEGRAM_CHAT_ID:
        raise ValueError("TELEGRAM_BOT_TOKEN или TELEGRAM_CHAT_ID не заданы в .env")
    resp = requests.post(
        f"https://api.telegram.org/bot{settings.TELEGRAM_BOT_TOKEN}/sendMessage",
        json={"chat_id": settings.TELEGRAM_CHAT_ID, "text": text},
        timeout=15,
    )
    data = resp.json()
    if not data.get("ok"):
        raise ValueError(data.get("description", "неизвестная ошибка Telegram"))
    return data


def telegram_send_today(request):
    if request.method != "POST":
        return JsonResponse({"error": "method not allowed"}, status=405)

    log = []
    try:
        _log(log, "Собираю задачи на сегодня...")
        text = _build_today_digest()
        _log(log, "Отправляю сообщение в Telegram...")
        send_telegram_message(text)
        _log(log, "Отправлено")
        logger.info("Отправлен дайджест в Telegram: %s", text.replace("\n", " | "))
        return JsonResponse({"ok": True, "log": log})
    except ValueError as e:
        _log(log, f"Ошибка: {e}")
        logger.warning("Не удалось отправить Telegram-дайджест: %s", e)
        return JsonResponse({"error": str(e), "log": log}, status=400)
    except requests.RequestException as e:
        _log(log, f"Сетевая ошибка: {e}")
        return JsonResponse({"error": "Не удалось связаться с Telegram", "log": log}, status=502)


# ---------- Чек-лист магазина ----------

def add_checklist_item(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST":
        text = request.POST.get("text", "").strip()
        if text:
            last_order = store.checklist_items.count()
            ChecklistItem.objects.create(store=store, text=text, order=last_order)
    return redirect("store_detail", store_id=store.id)


def toggle_checklist_item(request, item_id):
    item = get_object_or_404(ChecklistItem, id=item_id)
    if request.method == "POST":
        item.is_checked = not item.is_checked
        item.save(update_fields=["is_checked"])
    return redirect("store_detail", store_id=item.store_id)


def delete_checklist_item(request, item_id):
    item = get_object_or_404(ChecklistItem, id=item_id)
    store_id = item.store_id
    if request.method == "POST":
        item.delete()
    return redirect("store_detail", store_id=store_id)


def fill_checklist_template(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    if request.method == "POST" and store.store_type == "opening" and not store.checklist_items.exists():
        for i, (category, text) in enumerate(OPENING_CHECKLIST_TEMPLATE):
            ChecklistItem.objects.create(store=store, text=text, category=category, order=i)
    return redirect("store_detail", store_id=store.id)


# ---------- Архив заметок ----------

def archive_note(request, note_id):
    note = get_object_or_404(Note, id=note_id)
    store_id = note.task.store_id if note.task_id else None
    if request.method == "POST":
        note.is_archived = True
        note.save(update_fields=["is_archived"])
    if request.POST.get("next") == "dashboard":
        return redirect("dashboard")
    if store_id:
        return redirect("store_detail", store_id=store_id)
    return redirect("notes")


def notes_archive_page(request):
    notes = Note.objects.filter(is_archived=True).select_related("task__store").prefetch_related("files").order_by("-created_at")
    return render(request, "board/notes_archive.html", {"notes": notes})


# ---------- Калькулятор зарплаты ----------

def salary_page(request):
    settings_obj = SalarySettings.get()
    bonuses = Bonus.objects.all()

    grouped = {}
    for b in bonuses:
        key = (b.year, b.month)
        grouped.setdefault(key, []).append(b)

    ndfl = settings_obj.ndfl_percent
    oklad_net = round(float(settings_obj.base_oklad) * (1 - float(ndfl) / 100))

    months = []
    for (year, month), items in sorted(grouped.items(), reverse=True):
        bonus_gross = sum(b.amount for b in items)
        bonus_net = round(bonus_gross * (1 - float(ndfl) / 100))
        months.append({
            "year": year,
            "month": month,
            "month_name": [
                "", "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
                "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
            ][month],
            "bonuses": items,
            "bonus_gross": bonus_gross,
            "bonus_net": bonus_net,
            "total_net": oklad_net + bonus_net,
        })

    return render(request, "board/salary.html", {
        "settings_obj": settings_obj,
        "oklad_net": oklad_net,
        "months": months,
        "month_names": [
            "", "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
            "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
        ],
    })


def update_oklad(request):
    settings_obj = SalarySettings.get()
    if request.method == "POST":
        try:
            settings_obj.base_oklad = int(request.POST.get("base_oklad", settings_obj.base_oklad))
            settings_obj.ndfl_percent = request.POST.get("ndfl_percent", settings_obj.ndfl_percent)
            settings_obj.save()
        except (TypeError, ValueError):
            pass
    return redirect("salary_page")


def add_bonus(request):
    if request.method == "POST":
        try:
            year = int(request.POST.get("year"))
            month = int(request.POST.get("month"))
            amount = int(request.POST.get("amount"))
        except (TypeError, ValueError):
            return redirect("salary_page")
        description = request.POST.get("description", "").strip()
        Bonus.objects.create(year=year, month=month, amount=amount, description=description)
    return redirect("salary_page")


def delete_bonus(request, bonus_id):
    bonus = get_object_or_404(Bonus, id=bonus_id)
    if request.method == "POST":
        bonus.delete()
    return redirect("salary_page")


# ---------- Реконструкции (учёт премий) ----------

def reconstructions_page(request):
    records = ReconstructionRecord.objects.prefetch_related("documents").all()
    return render(request, "board/reconstructions.html", {"records": records, "branches": BRANCH_CHOICES})


def add_reconstruction(request):
    if request.method == "POST":
        store_number = request.POST.get("store_number", "").strip()
        if store_number:
            record = ReconstructionRecord.objects.create(
                branch=request.POST.get("branch", ""),
                store_number=store_number,
                reconstruction_date=request.POST.get("reconstruction_date") or None,
                expected_amount=request.POST.get("expected_amount") or None,
                notes=request.POST.get("notes", "").strip(),
            )
            for f in request.FILES.getlist("files"):
                ReconstructionDocument.objects.create(record=record, file=f)
    return redirect("reconstructions_page")


def toggle_reconstruction_paid(request, record_id):
    record = get_object_or_404(ReconstructionRecord, id=record_id)
    if request.method == "POST":
        record.paid = not record.paid
        record.paid_date = datetime.date.today() if record.paid else None
        record.save(update_fields=["paid", "paid_date"])
    return redirect("reconstructions_page")


def delete_reconstruction(request, record_id):
    record = get_object_or_404(ReconstructionRecord, id=record_id)
    if request.method == "POST":
        record.delete()
    return redirect("reconstructions_page")


def add_reconstruction_doc(request, record_id):
    record = get_object_or_404(ReconstructionRecord, id=record_id)
    if request.method == "POST":
        for f in request.FILES.getlist("files"):
            ReconstructionDocument.objects.create(record=record, file=f)
    return redirect("reconstructions_page")


def delete_reconstruction_doc(request, doc_id):
    doc = get_object_or_404(ReconstructionDocument, id=doc_id)
    if request.method == "POST":
        doc.delete()
    return redirect("reconstructions_page")


DEFAULT_TD_ZONES = [
    "Место ДМ", "Кассы", "ПЧ", "МЛ", "КСО", "ТБ", "Офис на приемке", "АДМ", "ТД",
]


def budgets_page(request):
    archived = request.GET.get("archived") == "1"
    calculations = BudgetCalculation.objects.filter(is_archived=archived)
    contractors = Contractor.objects.all()
    return render(request, "board/budgets.html", {
        "calculations": calculations,
        "contractors": contractors,
        "archived": archived,
        "branches": BRANCH_CHOICES,
    })


def add_budget(request):
    if request.method == "POST":
        title = request.POST.get("title", "").strip()
        contractor = Contractor.objects.filter(id=request.POST.get("contractor_id")).first()
        if title and contractor:
            vat_rate = BudgetSettings.get().vat_rate
            calc = BudgetCalculation.objects.create(
                contractor=contractor,
                contractor_name=contractor.name,
                title=title,
                branch=request.POST.get("branch", ""),
                store_number=request.POST.get("store_number", "").strip(),
                vat_rate=vat_rate,
                equipment_rows=copy.deepcopy(contractor.equipment_rows),
                work_rows=copy.deepcopy(contractor.work_rows),
                td_zones=[{"label": z, "qty": 0} for z in DEFAULT_TD_ZONES],
            )
            return redirect("budget_detail", calc_id=calc.id)
    return redirect("budgets_page")


def budget_detail(request, calc_id):
    calc = get_object_or_404(BudgetCalculation, id=calc_id)
    contractors = list(Contractor.objects.all())
    initial_state = {
        "id": calc.id,
        "title": calc.title,
        "branch": calc.branch,
        "storeNumber": calc.store_number,
        "contractorId": calc.contractor_id,
        "contractorName": calc.contractor_name,
        "vatRate": float(calc.vat_rate),
        "premium": float(calc.premium),
        "equipmentRows": calc.equipment_rows,
        "workRows": calc.work_rows,
        "tdZones": calc.td_zones,
        "isArchived": calc.is_archived,
        "hasReconstruction": calc.reconstruction_id is not None,
        "contractors": [
            {"id": c.id, "name": c.name, "equipmentRows": c.equipment_rows, "workRows": c.work_rows}
            for c in contractors
        ],
    }
    return render(request, "board/budget_detail.html", {
        "calc": calc,
        "initial_state": initial_state,
        "branches": BRANCH_CHOICES,
    })


@require_POST
def budget_save(request, calc_id):
    calc = get_object_or_404(BudgetCalculation, id=calc_id)
    try:
        payload = json.loads(request.body.decode("utf-8"))
    except (ValueError, UnicodeDecodeError):
        return JsonResponse({"ok": False, "error": "bad json"}, status=400)

    contractor = Contractor.objects.filter(id=payload.get("contractorId")).first()
    calc.contractor = contractor
    calc.contractor_name = contractor.name if contractor else (payload.get("contractorName") or calc.contractor_name)
    calc.title = (payload.get("title") or calc.title).strip() or calc.title
    calc.branch = payload.get("branch", calc.branch)
    calc.store_number = payload.get("storeNumber", calc.store_number)
    calc.vat_rate = payload.get("vatRate", calc.vat_rate)
    calc.premium = payload.get("premium", calc.premium)
    calc.equipment_rows = payload.get("equipmentRows") or []
    calc.work_rows = payload.get("workRows") or []
    calc.td_zones = payload.get("tdZones") or []

    totals = compute_budget_totals(calc.equipment_rows, calc.work_rows, calc.td_zones, calc.vat_rate, calc.premium)
    calc.total_amount = totals["grand_total"]
    calc.save()
    return JsonResponse({"ok": True, "totals": totals})


def toggle_budget_archived(request, calc_id):
    calc = get_object_or_404(BudgetCalculation, id=calc_id)
    if request.method == "POST":
        calc.is_archived = not calc.is_archived
        calc.save(update_fields=["is_archived", "updated_at"])
    return redirect("budgets_page")


def delete_budget(request, calc_id):
    calc = get_object_or_404(BudgetCalculation, id=calc_id)
    if request.method == "POST":
        calc.delete()
    return redirect("budgets_page")


def budget_to_reconstruction(request, calc_id):
    calc = get_object_or_404(BudgetCalculation, id=calc_id)
    if request.method == "POST" and not calc.reconstruction_id:
        record = ReconstructionRecord.objects.create(
            branch=calc.branch,
            store_number=calc.store_number or calc.title,
            expected_amount=int(calc.premium),
            notes=f"Создано из расчёта бюджета «{calc.title}», подрядчик {calc.contractor_name}",
        )
        calc.reconstruction = record
        calc.is_archived = True
        calc.save(update_fields=["reconstruction", "is_archived", "updated_at"])
    return redirect("reconstructions_page")


def budget_contractors(request):
    contractors = list(Contractor.objects.all())
    vat_rate = BudgetSettings.get().vat_rate
    initial_state = {
        "contractorOrder": [c.key for c in contractors],
        "contractors": {
            c.key: {
                "name": c.name,
                "equipmentRows": c.equipment_rows,
                "workRows": c.work_rows,
            }
            for c in contractors
        },
        "vatRate": float(vat_rate),
    }
    return render(request, "board/budget_contractors.html", {"initial_state": initial_state})


@require_POST
def budget_contractors_save(request):
    try:
        payload = json.loads(request.body.decode("utf-8"))
    except (ValueError, UnicodeDecodeError):
        return JsonResponse({"ok": False, "error": "bad json"}, status=400)

    contractor_order = payload.get("contractorOrder") or []
    contractors = payload.get("contractors") or {}

    existing_keys = set(Contractor.objects.values_list("key", flat=True))
    payload_keys = set(contractors.keys())
    Contractor.objects.filter(key__in=existing_keys - payload_keys).delete()

    for index, key in enumerate(contractor_order):
        data = contractors.get(key)
        if not data:
            continue
        name = (data.get("name") or "").strip()
        if not name:
            continue
        Contractor.objects.update_or_create(
            key=key,
            defaults={
                "name": name,
                "order": index,
                "equipment_rows": data.get("equipmentRows") or [],
                "work_rows": data.get("workRows") or [],
            },
        )

    if payload.get("vatRate") is not None:
        settings_obj = BudgetSettings.get()
        settings_obj.vat_rate = payload["vatRate"]
        settings_obj.save(update_fields=["vat_rate"])

    return JsonResponse({"ok": True})


def _save_router_config_values(store, post_data, keys):
    values = {key: (post_data.get(key) or "").strip() for key in keys}
    obj, _ = StoreRouterConfig.objects.get_or_create(store=store)
    obj.values = values
    obj.save(update_fields=["values", "updated_at"])
    return values


@require_POST
def save_router_config(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    _, keys = _router_config_placeholders(store)
    _save_router_config_values(store, request.POST, keys)
    return redirect(reverse("store_detail", args=[store.id]) + "#router-config")


@require_POST
def download_router_config(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    _, keys = _router_config_placeholders(store)
    values = _save_router_config_values(store, request.POST, keys)

    content = RouterConfigTemplate.get().template_text
    for key in keys:
        content = content.replace(f"<{key}>", values.get(key, ""))

    response = HttpResponse(content, content_type="text/plain; charset=utf-8")
    response["Content-Disposition"] = 'attachment; filename="fgt_system.conf"'
    return response


@require_POST
def recalc_device_ips(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    router_config = StoreRouterConfig.objects.filter(store=store).first()
    network = (router_config.values.get("NETWORK") if router_config else None)
    device_values = compute_device_ips_for_network(network)
    if device_values:
        StoreDeviceIPConfig.objects.update_or_create(store=store, defaults={"values": device_values})
    return redirect(reverse("store_detail", args=[store.id]) + "#device-ips")


@require_POST
def save_device_ips(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    values = {key: (request.POST.get(key) or "").strip() for key in DEVICE_IP_LABELS.keys()}
    obj, _ = StoreDeviceIPConfig.objects.get_or_create(store=store)
    obj.values = values
    obj.save(update_fields=["values", "updated_at"])
    return redirect(reverse("store_detail", args=[store.id]) + "#device-ips")


@require_POST
def regenerate_ascn_rows(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    kso_count_raw = (request.POST.get("kso_count") or "").strip()
    if kso_count_raw:
        store.kso_count = kso_count_raw
        store.save(update_fields=["kso_count"])
    try:
        target_count = int(store.kso_count)
    except (TypeError, ValueError):
        target_count = 0
    existing_count = store.cash_registers.count()
    if target_count > existing_count:
        StoreCashRegister.objects.bulk_create([
            StoreCashRegister(store=store, order=i, number=str(KSO_NUMBER_BASE + i))
            for i in range(existing_count, target_count)
        ])
    return redirect(reverse("store_detail", args=[store.id]) + "#ascn-generator")


@require_POST
def save_ascn_rows(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    store.locality = (request.POST.get("locality") or "").strip()
    store.fsrar_id = (request.POST.get("fsrar_id") or "").strip()
    store.save(update_fields=["locality", "fsrar_id"])
    for reg in store.cash_registers.all():
        prefix = f"row-{reg.id}-"
        reg.number = (request.POST.get(prefix + "number") or "").strip()
        reg.loymax_data = (request.POST.get(prefix + "loymax_data") or "").strip()
        reg.sbp_terminal = (request.POST.get(prefix + "sbp_terminal") or "").strip()
        reg.sbp_link = (request.POST.get(prefix + "sbp_link") or "").strip()
        reg.save(update_fields=["number", "loymax_data", "sbp_terminal", "sbp_link"])
    return redirect(reverse("store_detail", args=[store.id]) + "#ascn-generator")


def router_config_template(request):
    template = RouterConfigTemplate.get()
    if request.method == "POST":
        template.template_text = request.POST.get("template_text", "")
        template.save(update_fields=["template_text", "updated_at"])
        return redirect("router_config_template")
    return render(request, "board/router_config_template.html", {"template": template})


def login_view(request):
    error = None
    next_url = request.POST.get("next") or request.GET.get("next") or "/"
    if request.method == "POST":
        password = request.POST.get("password", "")
        if settings.SITE_PASSWORD and password == settings.SITE_PASSWORD:
            request.session["site_authenticated"] = True
            return redirect(next_url)
        error = "Неверный пароль."
    return render(request, "board/login.html", {"error": error, "next": next_url})


def logout_view(request):
    request.session.pop("site_authenticated", None)
    return redirect("login")
